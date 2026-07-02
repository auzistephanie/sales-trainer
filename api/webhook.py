"""AI 面試教練 Bot — Vercel Webhook Handler (interview_trainer edition)"""
import sys, os, re, traceback
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

from flask import Flask, request, jsonify
from dotenv import load_dotenv
load_dotenv(Path(__file__).parent.parent / ".env")

import requests as req
from datetime import datetime, timedelta

import re as _re

from interview_trainer import (
    generate_scenario, evaluate_response, analyze_conversation,
    QUESTION_TYPES, INDUSTRIES, DIFFICULTY_LEVELS, MBTI_COACHING,
    get_daily_tip, parse_resume,
    generate_job_questions, generate_job_tips,
    generate_cover_letter_from_jd, generate_tailored_cv_content, build_cv_docx,
    clean_jd_text, extract_company_role,
    calculate_cv_health, format_cv_health_message,
    generate_salary_benchmark, parse_salary_input,
    generate_negotiate_response, generate_negotiate_summary, extract_negotiate_reply,
    generate_debrief,
    calculate_ats_score, format_ats_message,
)
from utils import (
    load_stats, save_stats,
    load_session, save_session, clear_session,
    load_profile, save_profile,
    load_setup_session, save_setup_session, clear_setup_session,
    load_cv_text, save_cv_text,
    load_jd_session, save_jd_session, clear_jd_session,
    load_jobs, save_jobs,
    load_addjob_session, save_addjob_session, clear_addjob_session,
    load_negotiate_log, save_negotiate_log, load_debrief_log, save_debrief_log,
    send_telegram, send_document, upload_to_drive, set_current_chat_id,
    _redis_get, _redis_set, _redis_del,
)

_URL_RE = _re.compile(r'https?://\S+', _re.IGNORECASE)
from mbti_checker import (
    MBTI_QUESTIONS, MBTI_QUICK_DESC,
    validate_mbti, format_mbti_list, calculate_mbti,
    mbti_question_keyboard, mbti_question_text, mbti_result_text,
)

app = Flask(__name__)
TOKEN = lambda: os.getenv("TELEGRAM_BOT_TOKEN")

FREE_SESSION_LIMIT = 5
UPGRADE_MSG = (
    "🎓 你已完成 {n} 次免費練習！\n\n"
    "升級 Premium 解鎖無限練習 + 詳細進度分析：\n"
    "👉 [加入等候名單](https://t.me/salestraineraubot)（暫定 $68/月）\n\n"
    "或者繼續試用，每日送 1 次額外練習 🎁"
)

# ── 14 種行業選單 ─────────────────────────────────────────────────
SETUP_INDUSTRY_LIST = [
    ("💰", "金融/投行",    "金融／投資銀行"),
    ("💻", "科技/IT",      "科技／軟件開發"),
    ("📣", "市場/廣告",    "市場營銷／廣告"),
    ("🧩", "管理諮詢",     "管理諮詢"),
    ("🛍️", "零售/酒店",   "零售／酒店管理"),
    ("🚀", "初創",         "初創公司"),
    ("🏥", "醫療/健康科技", "醫療／藥劑／健康科技"),
    ("⚖️", "法律/合規",   "法律／合規"),
    ("👥", "人力資源",     "人力資源／招聘"),
    ("📚", "教育/培訓",    "教育／培訓"),
    ("🚚", "物流/供應鏈",  "物流／供應鏈"),
    ("🏗️", "地產/建築",  "地產／建築工程"),
    ("🎬", "傳媒/創意",   "傳媒／娛樂／創意"),
    ("🏛️", "政府/NGO",   "政府／NGO／公共服務"),
]

# MBTI 16種
MBTI_LIST = [
    "INTJ","INTP","ENTJ","ENTP",
    "INFJ","INFP","ENFJ","ENFP",
    "ISTJ","ISFJ","ESTJ","ESFJ",
    "ISTP","ISFP","ESTP","ESFP",
]

# ── Job Tracker 常數 ──────────────────────────────────────────────
JOB_STATUSES = ["Applied", "Phone Screen", "1st Interview", "2nd Interview", "Offer", "Rejected"]
STATUS_EMOJI  = {
    "Applied":       "📝",
    "Phone Screen":  "📞",
    "1st Interview": "🤝",
    "2nd Interview": "🔁",
    "Offer":         "🎉",
    "Rejected":      "❌",
}

DEBRIEF_PROMPT = (
    "請描述面試過程。包括：\n"
    "- 問咗咩問題\n"
    "- 你點答\n"
    "- Interviewer 嘅反應\n"
    "- 你覺得咩位答得唔好"
)


def _record_negotiate_summary(session: dict, summary: str):
    """談判結束時記低結果：有連結 job 就存落 job["negotiate_log"]，冇就存落全局 log。"""
    entry = {
        "date":    datetime.now().strftime("%Y-%m-%d"),
        "rounds":  session.get("round_num", 0),
        "summary": summary,
    }
    job_id = session.get("job_id")
    if job_id:
        jobs = load_jobs()
        job  = next((j for j in jobs if j["id"] == job_id), None)
        if job:
            job.setdefault("negotiate_log", []).append(entry)
            save_jobs(jobs)
            return
    log = load_negotiate_log()
    log.append(entry)
    save_negotiate_log(log)


def _record_debrief_result(job_info: dict, result: str):
    """覆盤分析完記低結果：有連結 job 就存落 job["debrief_log"]，冇就存落全局 log。"""
    entry = {"date": datetime.now().strftime("%Y-%m-%d"), "result": result}
    if job_info:
        jobs = load_jobs()
        job  = next((j for j in jobs if j["id"] == job_info.get("id")), None)
        if job:
            job.setdefault("debrief_log", []).append(entry)
            save_jobs(jobs)
            return
    log = load_debrief_log()
    log.append(entry)
    save_debrief_log(log)


# ── 工具 ──────────────────────────────────────────────────────────

def answer_cb(cb_id, text=""):
    req.post(
        f"https://api.telegram.org/bot{TOKEN()}/answerCallbackQuery",
        json={"callback_query_id": cb_id, "text": text, "show_alert": False},
        timeout=5,
    )

def cmd(text, command):
    return text == command or text.startswith(command + " ") or text.startswith(command + "@")


# ── Stats ──────────────────────────────────────────────────────────

def record_score(qtype_name: str, score: int):
    data    = load_stats()
    scores  = data.setdefault("qtype_scores", {})
    history = scores.setdefault(qtype_name, [])
    history.append(score)
    scores[qtype_name] = history[-20:]
    data["total_sessions"] = data.get("total_sessions", 0) + 1

    today     = datetime.now().strftime("%Y-%m-%d")
    yesterday = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")
    s = data.setdefault("streak", {"last_date": "", "count": 0})
    if s["last_date"] == today:
        pass
    elif s["last_date"] == yesterday:
        s["count"] += 1; s["last_date"] = today
    else:
        s["count"] = 1; s["last_date"] = today

    daily_log = data.setdefault("daily_log", {})
    daily_log[today] = daily_log.get(today, 0) + 1
    cutoff = (datetime.now() - timedelta(days=60)).strftime("%Y-%m-%d")
    data["daily_log"] = {d: c for d, c in daily_log.items() if d >= cutoff}

    score_log = data.setdefault("score_log", [])
    score_log.append({"date": today, "qtype": qtype_name, "score": score})
    data["score_log"] = score_log[-200:]

    save_stats(data)


def check_free_limit() -> bool:
    return True  # 限制已停用（Stephanie 自用，2026-07-01）— 想重開俾其他用戶：復原返落面呢兩行
    # data  = load_stats()
    # total = data.get("total_sessions", 0)
    # today_bonus = data.get("today_bonus_date") == datetime.now().strftime("%Y-%m-%d")
    # return total < FREE_SESSION_LIMIT or today_bonus


def use_daily_bonus():
    data = load_stats()
    data["today_bonus_date"] = datetime.now().strftime("%Y-%m-%d")
    save_stats(data)


def handle_stats():
    data   = load_stats()
    scores = data.get("qtype_scores", {})
    total  = data.get("total_sessions", 0)
    streak = data.get("streak", {})

    if not scores:
        send_telegram("未有練習記錄，用 /practice 開始！")
        return

    ranked = sorted(
        [(n, sum(sc)/len(sc), len(sc)) for n, sc in scores.items() if sc],
        key=lambda x: x[1], reverse=True
    )
    profile = load_profile()
    header  = f"🎯 {profile.get('job_title','')}  🧠 {profile.get('mbti','')}" if profile else ""

    lines = ["📊 面試訓練進度"]
    if header.strip(): lines.append(header)
    lines.append(f"總練習：{total} 次  |  連續 {streak.get('count',0)} 日\n")
    lines.append("【各題型掌握度】")
    for name, avg, count in ranked:
        lines.append(f"{'█'*round(avg)}{'░'*(4-round(avg))}  {name}  {round(avg/4*100)}%（{count}次）")

    weak = [x for x in ranked if x[1] < 2.5]
    if weak:
        lines.append(f"\n⚠️ 薄弱項目：{', '.join(x[0] for x in weak[:3])}")
        lines.append("用 /drill 針對練習 👆")

    send_telegram("\n".join(lines), reply_markup={"inline_keyboard": [[
        {"text": "🎯 繼續練習", "callback_data": "practice_new"},
    ]]})


# ── Setup ────────────────────────────────────────────────────────

def send_industry_keyboard(intro=False):
    kb = []
    row = []
    for i, (icon, short, _) in enumerate(SETUP_INDUSTRY_LIST):
        row.append({"text": f"{icon} {short}", "callback_data": f"setup_ind_{i}"})
        if len(row) == 2:
            kb.append(row); row = []
    if row: kb.append(row)
    kb.append([{"text": "✏️ 自定行業", "callback_data": "setup_ind_custom"}])
    msg = "👋 歡迎！先設定背景令練習更貼近你。\n\n🏭 目標行業係？" if intro else "🏭 揀目標行業："
    send_telegram(msg, reply_markup={"inline_keyboard": kb})


def send_mbti_keyboard():
    kb = []
    for i in range(0, 16, 4):
        kb.append([{"text": t, "callback_data": f"setup_mbti_{t}"} for t in MBTI_LIST[i:i+4]])
    kb.append([{"text": "⏭️ 唔知 / 跳過", "callback_data": "setup_mbti_skip"}])
    send_telegram(
        "🧠 你嘅 MBTI？（唔知可跳過，之後 /setup 再改）\n"
        "唔識做測試？👉 https://www.16personalities.com/ch",
        reply_markup={"inline_keyboard": kb},
    )


def send_setup_done(profile, extra=""):
    send_telegram(
        f"✅ 設定完成！\n\n"
        f"🎯 目標職位：{profile.get('job_title','未設定')}\n"
        f"🏭 行業：{profile.get('industry','未設定')}\n"
        f"🧠 MBTI：{profile.get('mbti','未設定')}"
        f"{extra}\n\n"
        f"開始你嘅面試練習！",
        reply_markup={"inline_keyboard": [[{"text": "🎯 立即練習", "callback_data": "practice_new"}]]},
    )


# ── Practice ──────────────────────────────────────────────────────

def start_practice(force_qtype=None, force_industry=None, difficulty=None):
    if not check_free_limit():
        data  = load_stats()
        total = data.get("total_sessions", 0)
        send_telegram(
            UPGRADE_MSG.format(n=total),
            reply_markup={"inline_keyboard": [[
                {"text": "🎁 今日免費練習", "callback_data": "claim_bonus"},
                {"text": "📊 睇進度",       "callback_data": "show_stats"},
            ]]}
        )
        return

    profile = load_profile() or {}
    if not force_industry:
        force_industry = profile.get("industry")

    display, scenario = generate_scenario(force_qtype, force_industry, difficulty)
    save_session({"state": "waiting_response", "scenario": scenario})
    send_telegram(display)


def handle_user_response(user_text):
    session = load_session()
    if not session or session.get("state") != "waiting_response":
        send_telegram("唔係練習模式。用 /practice 開始！")
        return

    clear_session()
    scenario   = session["scenario"]
    qtype_name = scenario["qtype"]["name"]

    send_telegram("🤔 AI 評估緊你嘅回答，稍等⋯⋯")
    feedback = evaluate_response(user_text, scenario, profile=load_profile() or {})

    match = re.search(r"評分[：:]\s*([1-4])", feedback)
    if match:
        record_score(qtype_name, int(match.group(1)))

    send_telegram(feedback, reply_markup={"inline_keyboard": [
        [
            {"text": "🔄 再練一個",              "callback_data": "practice_new"},
            {"text": f"🎯 再練「{qtype_name}」", "callback_data": f"drill_{qtype_name}"},
        ],
        [
            {"text": "📊 睇進度",  "callback_data": "show_stats"},
            {"text": "💡 今日技巧", "callback_data": "show_tip"},
        ],
    ]})


# ── MBTI Check Session ────────────────────────────────────────────
_MBTI_KEY = "mbti_check_session"

def load_mbti_session():   return _redis_get(_MBTI_KEY) or {}
def save_mbti_session(d):  _redis_set(_MBTI_KEY, d, ex=900)   # 15 min TTL
def clear_mbti_session():  _redis_del(_MBTI_KEY)


def start_mbti_check():
    """開始 MBTI 20題檢測。"""
    save_mbti_session({"step": 0, "answers": []})
    send_telegram(
        "🔍 *MBTI 檢測（20條問題，約5分鐘）*\n\n"
        "每個維度 5 條問題，參考 16personalities 設計。\n"
        "直接撳按鈕揀 A 或 B 就得！\n\n"
        f"{mbti_question_text(0)}",
        reply_markup=mbti_question_keyboard(0),
    )


def handle_mbti_answer(step: int, answer: str):
    """處理一個 MBTI 答案，繼續或顯示結果。"""
    sess    = load_mbti_session()
    answers = sess.get("answers", [])

    # 防止重複或亂序
    if len(answers) != step:
        send_telegram("⚠️ 出咗問題，用 /mbti 重新開始。")
        clear_mbti_session()
        return

    answers.append(answer)
    next_step = step + 1

    if next_step >= len(MBTI_QUESTIONS):
        # 完成 — 計算結果
        clear_mbti_session()
        result = calculate_mbti(answers)
        if not result:
            send_telegram("❌ 計算出錯，請用 /mbti 重新開始。")
            return

        # 儲存 MBTI 入 profile
        profile = load_profile()
        profile["mbti"] = result["mbti"]
        save_profile(profile)

        send_telegram(
            mbti_result_text(result),
            reply_markup={"inline_keyboard": [[
                {"text": "🎯 開始面試練習", "callback_data": "practice_new"},
                {"text": "⚙️ 睇我嘅設定",  "callback_data": "show_mystatus"},
            ]]}
        )
    else:
        save_mbti_session({"step": next_step, "answers": answers})
        send_telegram(
            mbti_question_text(next_step),
            reply_markup=mbti_question_keyboard(next_step),
        )


def cmd_mbti(text: str):
    """處理 /mbti 指令。"""
    parts = text.strip().split()
    if len(parts) > 1:
        # 直接輸入：/mbti INTJ
        mbti = parts[1].upper()
        if validate_mbti(mbti):
            profile = load_profile()
            profile["mbti"] = mbti
            save_profile(profile)
            coaching = MBTI_COACHING.get(mbti, {})
            note = f"\n\n💡 面試盲點：{coaching['watch_out']}" if coaching else ""
            send_telegram(
                f"✅ 已儲存 MBTI：*{mbti}* — {MBTI_QUICK_DESC.get(mbti, '')}{note}",
                reply_markup={"inline_keyboard": [[
                    {"text": "🎯 開始練習", "callback_data": "practice_new"},
                ]]}
            )
        else:
            send_telegram(
                f"❌ 唔識 `{mbti}`，請輸入以下 16 種之一：\n\n{format_mbti_list()}\n\n"
                f"例如：`/mbti INTJ`\n\n或者打 `/mbti` 開始 20 題檢測。"
            )
    else:
        # 開始檢測
        send_telegram(
            "📝 *點樣設定 MBTI？*\n\n"
            "*方法 1：做20題快速檢測*（推薦，約5分鐘）\n"
            "*方法 2：直接輸入*（如果你已知自己嘅 MBTI）\n"
            "例如：`/mbti INTJ`\n\n"
            "*方法 3：做官方測試*\n"
            "去 [16personalities.com](https://www.16personalities.com/ch)（約10分鐘）\n"
            "做完後用方法 2 輸入結果",
            reply_markup={"inline_keyboard": [
                [{"text": "🔍 開始20題檢測", "callback_data": "mbti_start"}],
                [{"text": "🌐 去 16personalities", "url": "https://www.16personalities.com/ch"}],
            ]}
        )


# ── Resume Parser ────────────────────────────────────────────────

def handle_document(document: dict):
    """處理用戶上傳嘅 resume（PDF / DOCX）。"""
    mime  = document.get("mime_type", "")
    fname = document.get("file_name", "").lower()

    is_pdf  = mime == "application/pdf" or fname.endswith(".pdf")
    is_docx = mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" \
              or fname.endswith(".docx")

    if not is_pdf and not is_docx:
        send_telegram("⚠️ 請上傳 PDF 或 Word (.docx) 格式嘅 resume。")
        return

    send_telegram("📄 收到 resume！AI 分析緊，稍等⋯⋯")

    # 1. 取得下載 URL
    file_info = req.get(
        f"https://api.telegram.org/bot{TOKEN()}/getFile",
        params={"file_id": document["file_id"]}, timeout=10,
    ).json()
    if not file_info.get("ok"):
        send_telegram("❌ 下載失敗，請重試。")
        return

    file_url  = f"https://api.telegram.org/file/bot{TOKEN()}/{file_info['result']['file_path']}"
    file_bytes = req.get(file_url, timeout=30).content

    # 2. 提取文字
    resume_text = ""
    try:
        if is_pdf:
            import pdfplumber, io
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                resume_text = "\n".join(p.extract_text() or "" for p in pdf.pages)
        else:
            import docx, io
            doc = docx.Document(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = "  |  ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        parts.append(row_text)
            resume_text = "\n".join(parts)
    except Exception as e:
        send_telegram(f"❌ 解析文件失敗：{e}\n請確保唔係掃描版圖片 PDF。")
        return

    if not resume_text.strip():
        send_telegram("⚠️ 未能提取文字，可能係掃描版 PDF。請試用文字版 PDF 或 .docx。")
        return

    # 3. AI 分析
    parsed = parse_resume(resume_text)
    if not parsed:
        send_telegram("❌ AI 分析失敗，請重試。")
        return

    # 4. 儲存 CV 全文 + 更新 profile + CV Health Score
    save_cv_text(resume_text)
    profile = load_profile() or {}
    if parsed.get("job_title"):       profile["job_title"] = parsed["job_title"]
    if parsed.get("industry"):        profile["industry"]  = parsed["industry"]
    if parsed.get("exp_years"):       profile["exp_years"] = str(parsed["exp_years"])
    if parsed.get("current_company"): profile["company"]   = parsed["current_company"]
    if parsed.get("key_skills"):      profile["key_skills"] = parsed["key_skills"]
    if parsed.get("education"):       profile["education"]  = parsed["education"]

    health = calculate_cv_health(resume_text)
    profile["cv_health_score"] = health["total"]
    save_profile(profile)

    # 5. 回覆確認
    lines = [
        "✅ *Resume 分析完成！*\n",
        f"📝 {parsed.get('summary', '')}",
        "",
        f"🎯 職位：{parsed.get('job_title', '未識別')}",
        f"🏭 行業：{parsed.get('industry', '未識別')}",
        f"⏱️ 年資：{parsed.get('exp_years', '?')} 年",
    ]
    if parsed.get("current_company"): lines.append(f"🏢 公司：{parsed['current_company']}")
    if parsed.get("key_skills"):      lines.append(f"🛠️ 技能：{parsed['key_skills']}")
    if parsed.get("education"):       lines.append(f"🎓 學歷：{parsed['education']}")

    setup = load_setup_session()
    onboarding = bool(setup) and setup.get("state") == "setup_cv_upload"

    if not onboarding:
        lines.append("\n面試練習會根據你嘅背景個人化！")
    send_telegram("\n".join(lines))
    send_telegram(format_cv_health_message(health))

    # 6. Onboarding 中 → 繼續落 Salary Benchmark 步；獨立上傳 → 顯示練習／設定按鈕
    if onboarding:
        save_setup_session({"state": "setup_salary"})
        send_telegram("💰 仲有最後一步！你目標月薪期望大概係幾多？（例如：38000 或 38K）")
    else:
        send_telegram(
            "繼續做咩？",
            reply_markup={"inline_keyboard": [[
                {"text": "🎯 立即練習", "callback_data": "practice_new"},
                {"text": "⚙️ 修改設定",  "callback_data": "setup_start"},
            ]]}
        )


# ── Drill ─────────────────────────────────────────────────────────
def handle_drill_menu():
    kb = []
    row = []
    for q in QUESTION_TYPES:
        row.append({"text": q["name"], "callback_data": f"drill_{q['name']}"})
        if len(row) == 2:
            kb.append(row); row = []
    if row: kb.append(row)
    kb.append([{"text": "🎲 隨機場景", "callback_data": "practice_new"}])
    send_telegram("🎯 揀你想針對練習的題型：", reply_markup={"inline_keyboard": kb})


# ── Job Tracker Handlers ──────────────────────────────────────────

def _save_addjob_final(addjob: dict):
    """Save completed addjob session as a job record."""
    import uuid
    jobs = load_jobs()
    job  = {
        "id":           str(uuid.uuid4())[:8],
        "company":      addjob.get("company", ""),
        "role":         addjob.get("role", ""),
        "jd":           addjob.get("jd", ""),
        "link":         addjob.get("link", ""),
        "applied_date": datetime.now().strftime("%Y-%m-%d"),
        "last_touch":   datetime.now().strftime("%Y-%m-%d"),
        "status":       "Applied",
    }
    jobs.append(job)
    save_jobs(jobs)
    clear_addjob_session()
    send_telegram(
        f"✅ *求職記錄已儲存！*\n\n"
        f"📝 {job['company']} — {job['role']}\n"
        f"申請日：{job['applied_date']}",
        reply_markup={"inline_keyboard": [
            [
                {"text": "❓ 睇面試問題", "callback_data": f"job_q_{job['id']}"},
                {"text": "💡 Key Tips",   "callback_data": f"job_tips_{job['id']}"},
            ],
            [{"text": "📋 睇所有申請", "callback_data": "show_listjobs"}],
        ]}
    )


def handle_addjob_start():
    save_addjob_session({"state": "addjob_company"})
    send_telegram(
        "📋 *新增求職記錄*\n\n"
        "第 1 步：公司名係？",
        reply_markup={"inline_keyboard": [[{"text": "❌ 取消", "callback_data": "addjob_cancel"}]]}
    )


def handle_listjobs():
    jobs = load_jobs()
    if not jobs:
        send_telegram(
            "未有求職記錄。\n用 /addjob 新增第一個！",
            reply_markup={"inline_keyboard": [[{"text": "➕ 新增", "callback_data": "addjob_start"}]]}
        )
        return
    for job in jobs:
        emoji = STATUS_EMOJI.get(job.get("status", "Applied"), "📝")
        lines = [
            f"{emoji} *{job['company']} — {job['role']}*",
            f"狀態：{job.get('status','Applied')}  |  申請日：{job.get('applied_date','-')}",
        ]
        if job.get("link"):
            lines.append(f"🔗 {job['link']}")
        kb = [
            [
                {"text": "❓ 面試問題", "callback_data": f"job_q_{job['id']}"},
                {"text": "💡 Key Tips",  "callback_data": f"job_tips_{job['id']}"},
            ],
            [
                {"text": "🎯 練習",     "callback_data": f"job_practice_{job['id']}"},
                {"text": "📊 更新狀態", "callback_data": f"job_updatestatus_{job['id']}"},
            ],
            [
                {"text": "📄 Cover Letter", "callback_data": f"job_cl_{job['id']}"},
                {"text": "📋 Tailored CV",  "callback_data": f"job_cv_{job['id']}"},
            ],
            [
                {"text": "🤝 Negotiate", "callback_data": f"job_negotiate_{job['id']}"},
            ],
        ]
        send_telegram("\n".join(lines), reply_markup={"inline_keyboard": kb})


def handle_job_questions(job_id: str):
    jobs = load_jobs()
    job  = next((j for j in jobs if j["id"] == job_id), None)
    if not job:
        send_telegram("⚠️ 搵唔到呢個申請記錄。")
        return
    send_telegram(f"🤔 AI 生成 *{job['company']}* 面試問題緊⋯⋯")
    result = generate_job_questions(job)
    send_telegram(result, reply_markup={"inline_keyboard": [[
        {"text": "🎯 針對練習", "callback_data": f"job_practice_{job_id}"},
    ]]})


def handle_job_tips(job_id: str):
    jobs = load_jobs()
    job  = next((j for j in jobs if j["id"] == job_id), None)
    if not job:
        send_telegram("⚠️ 搵唔到呢個申請記錄。")
        return
    send_telegram(f"💡 AI 生成 *{job['company']}* Key Talking Points⋯⋯")
    result = generate_job_tips(job)
    send_telegram(result, reply_markup={"inline_keyboard": [[
        {"text": "❓ 睇面試問題", "callback_data": f"job_q_{job_id}"},
        {"text": "🎯 練習",       "callback_data": f"job_practice_{job_id}"},
    ]]})


def handle_job_cover_letter(job_id: str):
    jobs = load_jobs()
    job  = next((j for j in jobs if j["id"] == job_id), None)
    if not job:
        send_telegram("⚠️ 搵唔到呢個申請記錄。")
        return
    cv_text = load_cv_text()
    if not cv_text:
        send_telegram("⚠️ 未有你嘅 CV 記錄。請先上傳 CV（PDF/.docx）到 bot。")
        return
    send_telegram(f"✍️ 根據你嘅 CV 生成 *{job['company']}* Cover Letter⋯⋯")
    result = generate_cover_letter_from_jd(cv_text, job.get("jd",""), job["company"], job["role"])
    send_telegram(
        f"📄 *Cover Letter — {job['company']}*\n_{job['role']}_\n\n{result}",
        reply_markup={"inline_keyboard": [[
            {"text": "📋 生成 Tailored CV", "callback_data": f"job_cv_{job_id}"},
            {"text": "🔄 重新生成",          "callback_data": f"job_cl_{job_id}"},
        ]]}
    )


def handle_job_tailored_cv(job_id: str):
    jobs = load_jobs()
    job  = next((j for j in jobs if j["id"] == job_id), None)
    if not job:
        send_telegram("⚠️ 搵唔到呢個申請記錄。")
        return
    cv_text = load_cv_text()
    if not cv_text:
        send_telegram("⚠️ 未有你嘅 CV 記錄。請先上傳 CV（PDF/.docx）到 bot。")
        return
    send_telegram(f"🛠️ 生成 *{job['company']}* Tailored CV 緊，約 15 秒⋯⋯")
    cv_data = generate_tailored_cv_content(cv_text, job.get("jd",""), job["company"], job["role"])
    if not cv_data:
        send_telegram("❌ 生成失敗，請重試。")
        return
    try:
        docx_bytes = build_cv_docx(cv_data, job["company"], job["role"])
        filename   = f"CV_{job['company'].replace(' ','_')[:20]}_{job['role'].replace(' ','_')[:15]}.docx"
        drive_link = upload_to_drive(docx_bytes, filename)
        if drive_link:
            send_telegram(
                f"✅ Tailored CV 已上傳 Google Drive！\n\n"
                f"📋 *{job['role']} @ {job['company']}*\n"
                f"[🔗 打開 CV]({drive_link})",
                reply_markup={"inline_keyboard": [[
                    {"text": "📄 生成 Cover Letter", "callback_data": f"job_cl_{job_id}"},
                ]]}
            )
        else:
            # Fallback: send file directly via Telegram
            send_document(docx_bytes, filename, caption=f"📋 Tailored CV for {job['role']} @ {job['company']}")
            send_telegram(
                "✅ Tailored CV 已生成！（Drive 未設定，直接傳送）",
                reply_markup={"inline_keyboard": [[
                    {"text": "📄 生成 Cover Letter", "callback_data": f"job_cl_{job_id}"},
                ]]}
            )
        ats = calculate_ats_score(job.get("jd", ""), cv_text)
        profile = load_profile() or {}
        send_telegram(format_ats_message(ats, profile.get("cv_health_score")))
        job["ats_score"] = ats.get("score")
        if drive_link:
            job["cv_drive_link"] = drive_link
        save_jobs(jobs)
    except Exception as e:
        send_telegram(f"❌ 生成 .docx 失敗：{e}")


def handle_job_questions_from_jd(sess: dict):
    fake_job = {"company": sess.get("company",""), "role": sess.get("role",""), "jd": sess.get("jd_text","")}
    send_telegram(f"🤔 生成 *{fake_job['company']}* 面試問題⋯⋯")
    send_telegram(generate_job_questions(fake_job), reply_markup={"inline_keyboard": [[
        {"text": "📄 生成 Cover Letter", "callback_data": "jd_cover_letter"},
    ]]})

def handle_job_tips_from_jd(sess: dict):
    fake_job = {"company": sess.get("company",""), "role": sess.get("role",""), "jd": sess.get("jd_text","")}
    send_telegram(f"💡 生成 *{fake_job['company']}* Key Tips⋯⋯")
    send_telegram(generate_job_tips(fake_job), reply_markup={"inline_keyboard": [[
        {"text": "📄 生成 Cover Letter", "callback_data": "jd_cover_letter"},
    ]]})


def handle_update_status_menu(job_id: str):
    jobs = load_jobs()
    job  = next((j for j in jobs if j["id"] == job_id), None)
    if not job:
        send_telegram("⚠️ 搵唔到呢個申請記錄。")
        return
    kb = [[{"text": f"{STATUS_EMOJI.get(s,'')} {s}", "callback_data": f"job_status_{job_id}_{s}"}]
          for s in JOB_STATUSES]
    send_telegram(
        f"📊 更新 *{job['company']} — {job['role']}* 狀態：",
        reply_markup={"inline_keyboard": kb}
    )


# ── JD / Cover Letter / Tailored CV Handlers ─────────────────────

def fetch_jd_via_jina(url: str) -> str:
    """用 Jina Reader 抓取 URL 內容（支援 JS 渲染頁面）。

    2026-07-02 修正：
    - 移除 X-No-Cache（之前強制每次完整重新 render JS 頁面，超慢 → timeout）
    - timeout 15→45（SEEK/JobsDB 冷 render 成日過 15s；Vercel maxDuration=60 有位）
    - 加 X-Return-Format=markdown（乾淨 JD，唔使成堆 HTML）
    - 有 JINA_API_KEY 就用（免費層 rate limit 好鬆），冇都照跑
    - 加一次 retry
    """
    jina_key = os.environ.get("JINA_API_KEY", "").strip()
    # 被 Cloudflare / 反爬擋嘅頁面特徵 —— 抓到呢啲當失敗
    block_signs = (
        "just a moment", "error 403", "403: forbidden", "attention required",
        "enable javascript and cookies", "checking your browser", "cf-chl",
        "returned error 4", "returned error 5",
    )
    for attempt in range(2):
        headers = {"Accept": "text/plain", "X-Return-Format": "markdown"}
        if jina_key:
            headers["Authorization"] = f"Bearer {jina_key}"
        # 第二次用 browser engine（有 key 先啟用），可過部分 Cloudflare 挑戰
        if attempt == 1 and jina_key:
            headers["X-Engine"] = "browser"
        try:
            resp = req.get(f"https://r.jina.ai/{url}", headers=headers, timeout=45)
            text = resp.text.strip()
            low  = text[:600].lower()
            if resp.ok and len(text) > 200 and not any(s in low for s in block_signs):
                return text[:4000]
            print(f"[jina fetch] attempt {attempt+1} status={resp.status_code} len={len(text)} blocked={any(s in low for s in block_signs)}")
        except Exception as e:
            print(f"[jina fetch] attempt {attempt+1} {e}")
    return ""


def handle_url_message(url: str):
    """用戶發咗一條 URL — 嘗試抓 JD，然後問佢要做咩。"""
    send_telegram("🔍 抓取職位資料中⋯⋯")
    jd_text = fetch_jd_via_jina(url)

    if not jd_text:
        # Jina 抓唔到，叫用戶貼文字
        save_jd_session({"state": "waiting_jd_text", "url": url})
        send_telegram(
            "⚠️ 未能自動抓取內容（可能係需要登入）\n\n"
            "請直接貼上 JD 文字，我幫你繼續：",
            reply_markup={"inline_keyboard": [[{"text": "❌ 取消", "callback_data": "jd_cancel"}]]}
        )
        return

    # 成功抓到 → 清走 Jina boilerplate，再用 DeepSeek 乾淨抽 company/role
    jd_text = clean_jd_text(jd_text)
    info = extract_company_role(jd_text)
    company_guess = info.get("company", "")
    role_guess    = info.get("role", "")

    # 抽唔到 company 同 role（例如被擋、內容係雜訊）→ 當抓唔到，叫貼 JD
    if not company_guess and not role_guess:
        save_jd_session({"state": "waiting_jd_text", "url": url})
        send_telegram(
            "⚠️ 呢個網站擋咗自動抓取（或者要登入）\n\n"
            "請直接貼上 JD 文字，我幫你繼續：",
            reply_markup={"inline_keyboard": [[{"text": "❌ 取消", "callback_data": "jd_cancel"}]]}
        )
        return

    save_jd_session({
        "state":   "jd_ready",
        "url":     url,
        "jd_text": jd_text,
        "company": company_guess,
        "role":    role_guess,
    })

    send_telegram(
        f"✅ 成功抓取職位資料！\n\n"
        f"🏢 可能係：{company_guess}\n"
        f"🎯 職位：{role_guess}\n\n"
        "幫你做咩？",
        reply_markup={"inline_keyboard": [
            [{"text": "📄 生成 Cover Letter", "callback_data": "jd_cover_letter"}],
            [{"text": "📋 生成 Tailored CV",  "callback_data": "jd_tailored_cv"}],
            [
                {"text": "❓ 面試問題", "callback_data": "jd_questions"},
                {"text": "💡 Key Tips",  "callback_data": "jd_tips"},
            ],
            [{"text": "➕ 加入申請追蹤", "callback_data": "jd_add_to_tracker"}],
            [{"text": "✏️ 修正公司/職位",  "callback_data": "jd_edit_info"}],
        ]}
    )


def handle_jd_cover_letter():
    """根據儲存嘅 JD session 生成 Cover Letter。"""
    sess = load_jd_session()
    if not sess:
        send_telegram("⚠️ 冇 JD 資料，請重新貼 link 或 JD 文字。")
        return
    cv_text = load_cv_text()
    if not cv_text:
        send_telegram(
            "⚠️ 未有你嘅 CV 記錄。請先上傳你嘅 CV（PDF 或 .docx），再生成 Cover Letter。"
        )
        return
    send_telegram("✍️ AI 根據你嘅 CV 生成 Cover Letter 緊，稍等⋯⋯")
    result = generate_cover_letter_from_jd(
        cv_text,
        sess.get("jd_text", ""),
        sess.get("company", ""),
        sess.get("role", ""),
    )
    send_telegram(
        f"📄 *Cover Letter — {sess.get('company','')}*\n_{sess.get('role','')}_\n\n{result}",
        reply_markup={"inline_keyboard": [
            [{"text": "📋 生成 Tailored CV",   "callback_data": "jd_tailored_cv"}],
            [{"text": "➕ 加入申請追蹤",        "callback_data": "jd_add_to_tracker"}],
            [{"text": "🔄 重新生成",            "callback_data": "jd_cover_letter"}],
        ]}
    )


def handle_jd_tailored_cv():
    """生成 Tailored CV .docx 並發送。"""
    sess = load_jd_session()
    if not sess:
        send_telegram("⚠️ 冇 JD 資料，請重新貼 link 或 JD 文字。")
        return
    cv_text = load_cv_text()
    if not cv_text:
        send_telegram("⚠️ 未有你嘅 CV 記錄。請先上傳 CV（PDF/.docx）到 bot。")
        return
    send_telegram("🛠️ AI 生成 Tailored CV 緊，稍等約 15 秒⋯⋯")
    company = sess.get("company", "")
    role    = sess.get("role", "")
    cv_data = generate_tailored_cv_content(cv_text, sess.get("jd_text", ""), company, role)
    if not cv_data:
        send_telegram("❌ 生成失敗，請重試。")
        return
    try:
        docx_bytes = build_cv_docx(cv_data, company, role)
        filename   = f"CV_Tailored_{company.replace(' ','_')[:20]}.docx"
        drive_link = upload_to_drive(docx_bytes, filename)
        if drive_link:
            send_telegram(
                f"✅ Tailored CV 已上傳 Google Drive！\n\n"
                f"📋 *{role} @ {company}*\n"
                f"[🔗 打開 CV]({drive_link})",
                reply_markup={"inline_keyboard": [
                    [{"text": "📄 生成 Cover Letter", "callback_data": "jd_cover_letter"}],
                    [{"text": "➕ 加入申請追蹤",      "callback_data": "jd_add_to_tracker"}],
                ]}
            )
        else:
            # Fallback: send file directly
            send_document(docx_bytes, filename, caption=f"📄 Tailored CV — {role} @ {company}")
            send_telegram(
                "✅ Tailored CV 已生成！",
                reply_markup={"inline_keyboard": [
                    [{"text": "📄 生成 Cover Letter", "callback_data": "jd_cover_letter"}],
                    [{"text": "➕ 加入申請追蹤",      "callback_data": "jd_add_to_tracker"}],
                ]}
            )
        ats = calculate_ats_score(sess.get("jd_text", ""), cv_text)
        profile = load_profile() or {}
        send_telegram(format_ats_message(ats, profile.get("cv_health_score")))
        sess["ats_score"] = ats.get("score")
        if drive_link:
            sess["cv_drive_link"] = drive_link
        save_jd_session(sess)
    except Exception as e:
        send_telegram(f"❌ 生成 .docx 失敗：{e}")


def handle_jd_add_to_tracker():
    """將 JD session 嘅職位加入申請追蹤。"""
    import uuid
    sess = load_jd_session()
    if not sess:
        send_telegram("⚠️ 冇 JD 資料。")
        return
    jobs = load_jobs()
    job  = {
        "id":           str(uuid.uuid4())[:8],
        "company":      sess.get("company", ""),
        "role":         sess.get("role", ""),
        "jd":           sess.get("jd_text", "")[:500],
        "link":         sess.get("url", ""),
        "applied_date": datetime.now().strftime("%Y-%m-%d"),
        "last_touch":   datetime.now().strftime("%Y-%m-%d"),
        "status":       "Applied",
        "ats_score":      sess.get("ats_score"),
        "cv_drive_link":  sess.get("cv_drive_link"),
    }
    jobs.append(job)
    save_jobs(jobs)
    send_telegram(
        f"✅ 已加入追蹤！\n{job['company']} — {job['role']}\n申請日：{job['applied_date']}",
        reply_markup={"inline_keyboard": [[{"text": "📋 睇所有申請", "callback_data": "show_listjobs"}]]}
    )


# ── Callback ──────────────────────────────────────────────────────

def handle_callback(cb):
    answer_cb(cb["id"])
    data = cb.get("data", "")

    if data == "practice_new":
        start_practice()

    elif data == "claim_bonus":
        use_daily_bonus()
        start_practice()

    elif data == "show_stats":
        handle_stats()

    elif data == "show_tip":
        send_telegram(f"💡 今日面試技巧\n\n{get_daily_tip()}")

    elif data == "setup_start":
        clear_setup_session()
        send_industry_keyboard(intro=False)

    elif data == "setup_ind_custom":
        save_setup_session({"state": "setup_industry_custom"})
        send_telegram("✏️ 打出你嘅目標行業（例如：航空、網絡安全、遊戲）：")

    elif data.startswith("setup_ind_"):
        idx = int(data[len("setup_ind_"):])
        if idx < len(SETUP_INDUSTRY_LIST):
            _, _, full_name = SETUP_INDUSTRY_LIST[idx]
            profile = load_profile() or {}
            profile["industry"] = full_name
            save_profile(profile)
            save_setup_session({"state": "setup_jobtitle"})
            send_telegram(f"✅ 行業：{full_name}\n\n🎯 你嘅目標職位係？（例如：Product Manager）")

    elif data.startswith("setup_mbti_"):
        mbti_val = data[len("setup_mbti_"):]
        profile  = load_profile() or {}
        if mbti_val != "skip":
            profile["mbti"] = mbti_val
            save_profile(profile)
            extra = ""
            if mbti_val in MBTI_COACHING:
                extra = f"\n\n💡 {mbti_val} 面試盲點：{MBTI_COACHING[mbti_val]['watch_out']}"
            clear_setup_session()
            send_setup_done(profile, extra=extra)
        else:
            clear_setup_session()
            send_setup_done(profile)

    elif data == "mbti_start":
        start_mbti_check()

    elif data.startswith("mbti_ans_"):
        # mbti_ans_{step}_{A|B}
        parts = data.split("_")
        step_i = int(parts[2])
        ans    = parts[3]
        handle_mbti_answer(step_i, ans)

    elif data == "show_mystatus":
        # inline 版 mystatus
        profile = load_profile() or {}
        data2   = load_stats()
        total   = data2.get("total_sessions", 0)
        streak  = data2.get("streak", {}).get("count", 0)
        mbti    = profile.get("mbti", "未設定")
        note    = ""
        if profile.get("mbti") and profile["mbti"] in MBTI_COACHING:
            note = f"\n💡 盲點：{MBTI_COACHING[profile['mbti']]['watch_out']}"
        send_telegram(
            f"⚙️ 我的設定\n\n"
            f"🎯 目標職位：{profile.get('job_title','未設定')}\n"
            f"🏭 行業：{profile.get('industry','未設定')}\n"
            f"🧠 MBTI：{mbti}{note}\n\n"
            f"📊 總練習：{total} 次  🔥 連續：{streak} 日",
            reply_markup={"inline_keyboard": [[
                {"text": "🎯 開始練習", "callback_data": "practice_new"},
                {"text": "🔍 重做 MBTI", "callback_data": "mbti_start"},
            ]]}
        )

    elif data.startswith("drill_"):
        qtype_name = data[6:]
        start_practice(force_qtype=qtype_name)

    elif data == "onboard_cv":
        save_setup_session({"state": "setup_cv_upload"})
        send_telegram(
            "📄 請直接將你嘅 CV 拖入呢個 chat（PDF 或 Word .docx）。\n\n"
            "AI 會自動分析你嘅背景，然後開始個人化練習！"
        )

    elif data == "onboard_manual":
        clear_setup_session()
        send_industry_keyboard(intro=False)

    elif data == "review_start":
        save_session({"state": "waiting_review"})
        send_telegram(
            "📋 貼上你嘅真實面試問答記錄，AI 幫你分析失分點：\n\n"
            "（格式：面試官問：xxx\n我答：xxx）"
        )

    elif data == "jd_cover_letter":
        handle_jd_cover_letter()

    elif data == "jd_tailored_cv":
        handle_jd_tailored_cv()

    elif data == "jd_questions":
        sess = load_jd_session()
        if sess:
            handle_job_questions_from_jd(sess)
        else:
            send_telegram("⚠️ 冇 JD 資料，請重新貼 link。")

    elif data == "jd_tips":
        sess = load_jd_session()
        if sess:
            handle_job_tips_from_jd(sess)
        else:
            send_telegram("⚠️ 冇 JD 資料，請重新貼 link。")

    elif data == "jd_add_to_tracker":
        handle_jd_add_to_tracker()

    elif data == "jd_cancel":
        clear_jd_session()
        send_telegram("✅ 已取消。")

    elif data == "jd_edit_info":
        save_jd_session({**load_jd_session(), "state": "waiting_company"})
        send_telegram("✏️ 請輸入公司名稱：")

    elif data == "addjob_start":
        handle_addjob_start()

    elif data == "show_listjobs":
        handle_listjobs()

    elif data == "addjob_cancel":
        clear_addjob_session()
        send_telegram("✅ 已取消。用 /addjob 再試。")

    elif data == "addjob_skip_jd":
        addjob = load_addjob_session()
        if addjob:
            addjob["jd"]    = ""
            addjob["state"] = "addjob_link"
            save_addjob_session(addjob)
        send_telegram(
            "第 4 步：貼上職位連結（或 /skip 跳過）：",
            reply_markup={"inline_keyboard": [[{"text": "⏭️ 跳過", "callback_data": "addjob_skip_link"}]]}
        )

    elif data == "addjob_skip_link":
        addjob = load_addjob_session()
        if addjob:
            addjob["link"] = ""
            _save_addjob_final(addjob)

    elif data.startswith("job_cl_"):
        handle_job_cover_letter(data[7:])

    elif data.startswith("job_cv_"):
        handle_job_tailored_cv(data[7:])

    elif data.startswith("job_q_"):
        handle_job_questions(data[6:])

    elif data.startswith("job_tips_"):
        handle_job_tips(data[9:])

    elif data.startswith("job_practice_"):
        job_id = data[13:]
        jobs   = load_jobs()
        job    = next((j for j in jobs if j["id"] == job_id), None)
        if job:
            profile = load_profile() or {}
            if job.get("role"):
                profile["job_title"] = job["role"]
                save_profile(profile)
            start_practice(force_industry=profile.get("industry"))
        else:
            send_telegram("⚠️ 搵唔到呢個申請記錄。")

    elif data.startswith("job_updatestatus_"):
        handle_update_status_menu(data[17:])

    elif data.startswith("job_negotiate_"):
        job_id = data[len("job_negotiate_"):]
        jobs   = load_jobs()
        job    = next((j for j in jobs if j["id"] == job_id), None)
        if not job:
            send_telegram("⚠️ 搵唔到呢個申請記錄。")
        else:
            offer_details = f"職位：{job['role']}\n公司：{job['company']}\n（其他 package 詳情可以喺對話中補充）"
            save_session({"state": "negotiate_session", "offer_details": offer_details, "round_num": 0, "history": [], "job_id": job_id})
            send_telegram(
                f"🤝 即將同 *{job['company']}* 嘅 HR 傾 *{job['role']}* 嘅薪酬。\n\n打你想講嘅第一句：",
                reply_markup={"inline_keyboard": [[{"text": "🏁 結束談判", "callback_data": "negotiate_end"}]]}
            )

    elif data == "negotiate_cancel":
        clear_session()
        send_telegram("已取消談判練習。")

    elif data == "negotiate_end":
        session = load_session() or {}
        send_telegram("📊 生成談判總結⋯⋯")
        summary = generate_negotiate_summary(session.get("history", []))
        send_telegram(summary, reply_markup={"inline_keyboard": [[
            {"text": "🎯 繼續練習", "callback_data": "practice_new"},
        ]]})
        _record_negotiate_summary(session, summary)
        clear_session()

    elif data == "debrief_job_skip":
        save_session({"state": "debrief_input", "job_info": None})
        send_telegram(DEBRIEF_PROMPT)

    elif data.startswith("debrief_job_"):
        job_id = data[len("debrief_job_"):]
        jobs   = load_jobs()
        job    = next((j for j in jobs if j["id"] == job_id), None)
        save_session({"state": "debrief_input", "job_info": job})
        send_telegram(DEBRIEF_PROMPT)

    elif data.startswith("job_status_"):
        # job_status_{id}_{status} — status 可能包含空格，所以 split 限3份
        parts  = data.split("_", 3)   # ["job", "status", id, status_value]
        if len(parts) == 4:
            job_id     = parts[2]
            new_status = parts[3]
            jobs       = load_jobs()
            today      = datetime.now().strftime("%Y-%m-%d")
            old_status = None
            for j in jobs:
                if j["id"] == job_id:
                    old_status     = j.get("status")
                    j["status"]     = new_status
                    j["last_touch"] = today
                    j.pop("snooze_until", None)
                    break
            save_jobs(jobs)
            job = next((j for j in jobs if j["id"] == job_id), {})
            if old_status is not None and old_status != new_status:
                stats_data = load_stats()
                log = stats_data.setdefault("status_change_log", [])
                log.append({
                    "date": today, "company": job.get("company", ""),
                    "role": job.get("role", ""), "from": old_status, "to": new_status,
                })
                stats_data["status_change_log"] = log[-200:]
                save_stats(stats_data)
            emoji = STATUS_EMOJI.get(new_status, "📝")
            send_telegram(
                f"{emoji} 已更新：*{job.get('company','')} — {job.get('role','')}*\n"
                f"狀態 → {new_status}"
            )

    elif data.startswith("job_followup_"):
        job_id = data[len("job_followup_"):]
        jobs   = load_jobs()
        today  = datetime.now().strftime("%Y-%m-%d")
        for j in jobs:
            if j["id"] == job_id:
                j["last_touch"] = today
                j.pop("snooze_until", None)
                break
        save_jobs(jobs)
        job = next((j for j in jobs if j["id"] == job_id), {})
        send_telegram(f"✅ 已記錄 follow-up：{job.get('company','')} — {job.get('role','')}")

    elif data.startswith("job_snooze_"):
        job_id      = data[len("job_snooze_"):]
        jobs        = load_jobs()
        snooze_date = (datetime.now() + timedelta(days=3)).strftime("%Y-%m-%d")
        for j in jobs:
            if j["id"] == job_id:
                j["snooze_until"] = snooze_date
                break
        save_jobs(jobs)
        job = next((j for j in jobs if j["id"] == job_id), {})
        send_telegram(f"⏰ 已 snooze 3 日：{job.get('company','')} — {job.get('role','')}（{snooze_date} 再提你）")


# ── Message ───────────────────────────────────────────────────────

def handle_message(text):
    # JD session flow（waiting for JD text after failed link fetch）
    jd_sess = load_jd_session()
    if jd_sess and not text.startswith("/"):
        state = jd_sess.get("state", "")

        if state == "waiting_jd_text":
            jd_sess["jd_text"] = text.strip()
            jd_sess["state"]   = "jd_ready"
            save_jd_session(jd_sess)
            send_telegram(
                "✅ JD 已收到！幫你做咩？",
                reply_markup={"inline_keyboard": [
                    [{"text": "📄 生成 Cover Letter", "callback_data": "jd_cover_letter"}],
                    [{"text": "📋 生成 Tailored CV",  "callback_data": "jd_tailored_cv"}],
                    [
                        {"text": "❓ 面試問題", "callback_data": "jd_questions"},
                        {"text": "💡 Key Tips",  "callback_data": "jd_tips"},
                    ],
                    [{"text": "➕ 加入申請追蹤", "callback_data": "jd_add_to_tracker"}],
                ]}
            )
            return

        if state == "waiting_company":
            jd_sess["company"] = text.strip()
            jd_sess["state"]   = "waiting_role"
            save_jd_session(jd_sess)
            send_telegram("✏️ 職位名稱係？")
            return

        if state == "waiting_role":
            jd_sess["role"]  = text.strip()
            jd_sess["state"] = "jd_ready"
            save_jd_session(jd_sess)
            send_telegram(
                f"✅ {jd_sess.get('company','')} — {jd_sess['role']}\n\n幫你做咩？",
                reply_markup={"inline_keyboard": [
                    [{"text": "📄 生成 Cover Letter", "callback_data": "jd_cover_letter"}],
                    [{"text": "📋 生成 Tailored CV",  "callback_data": "jd_tailored_cv"}],
                    [{"text": "➕ 加入申請追蹤",      "callback_data": "jd_add_to_tracker"}],
                ]}
            )
            return

    # URL 偵測（唔係 session 狀態，直接貼 link）
    url_match = _URL_RE.search(text)
    if url_match and not text.startswith("/"):
        handle_url_message(url_match.group(0))
        return

    # AddJob flow（優先於 setup flow）
    addjob = load_addjob_session()
    if addjob and not text.startswith("/"):
        state = addjob.get("state")

        if state == "addjob_company":
            addjob["company"] = text.strip()
            addjob["state"]   = "addjob_role"
            save_addjob_session(addjob)
            send_telegram(
                f"✅ 公司：{text.strip()}\n\n第 2 步：職位名稱係？",
                reply_markup={"inline_keyboard": [[{"text": "❌ 取消", "callback_data": "addjob_cancel"}]]}
            )
            return

        if state == "addjob_role":
            addjob["role"]  = text.strip()
            addjob["state"] = "addjob_jd"
            save_addjob_session(addjob)
            send_telegram(
                f"✅ 職位：{text.strip()}\n\n第 3 步：貼上 JD 描述（或輸入 /skip 跳過）：",
                reply_markup={"inline_keyboard": [[{"text": "⏭️ 跳過", "callback_data": "addjob_skip_jd"}]]}
            )
            return

        if state == "addjob_jd":
            addjob["jd"]    = text.strip()
            addjob["state"] = "addjob_link"
            save_addjob_session(addjob)
            send_telegram(
                "✅ JD 已儲存！\n\n第 4 步：貼上職位連結（或輸入 /skip 跳過）：",
                reply_markup={"inline_keyboard": [[{"text": "⏭️ 跳過", "callback_data": "addjob_skip_link"}]]}
            )
            return

        if state == "addjob_link":
            addjob["link"] = text.strip()
            _save_addjob_final(addjob)
            return

    # AddJob skip callbacks via text fallback
    if text == "/skip":
        addjob = load_addjob_session()
        if addjob:
            state = addjob.get("state")
            if state == "addjob_jd":
                addjob["jd"]    = ""
                addjob["state"] = "addjob_link"
                save_addjob_session(addjob)
                send_telegram(
                    "第 4 步：貼上職位連結（或 /skip 跳過）：",
                    reply_markup={"inline_keyboard": [[{"text": "⏭️ 跳過", "callback_data": "addjob_skip_link"}]]}
                )
                return
            if state == "addjob_link":
                addjob["link"] = ""
                _save_addjob_final(addjob)
                return

    # Setup flow
    setup = load_setup_session()
    if setup and not text.startswith("/"):
        state = setup.get("state")

        if state == "setup_industry_custom":
            profile = load_profile() or {}
            profile["industry"] = text.strip()
            save_profile(profile)
            save_setup_session({"state": "setup_jobtitle"})
            send_telegram(f"✅ 行業：{text.strip()}\n\n🎯 你嘅目標職位係？")
            return

        if state == "setup_jobtitle":
            profile = load_profile() or {}
            profile["job_title"] = text.strip()
            save_profile(profile)
            save_setup_session({"state": "setup_salary"})
            send_telegram("💰 仲有最後一步！你目標月薪期望大概係幾多？（例如：38000 或 38K）")
            return

        if state == "setup_salary":
            profile = load_profile() or {}
            expected_salary = parse_salary_input(text)
            profile["expected_salary"]  = expected_salary
            profile["salary_currency"]  = "HKD"
            save_profile(profile)
            send_telegram("💰 分析緊薪酬市場數據⋯⋯")
            benchmark = generate_salary_benchmark(
                profile.get("job_title", "未指定職位"), expected_salary, profile.get("industry", "")
            )
            send_telegram(f"💰 HK 市場薪酬參考（2026）\n\n{benchmark}")
            save_setup_session({"state": "setup_mbti"})
            send_mbti_keyboard()
            return

        if state == "setup_mbti":
            mbti_input = text.strip().upper()
            if mbti_input in MBTI_COACHING:
                profile = load_profile() or {}
                profile["mbti"] = mbti_input
                save_profile(profile)
                clear_setup_session()
                extra = f"\n\n💡 {mbti_input} 面試盲點：{MBTI_COACHING[mbti_input]['watch_out']}"
                send_setup_done(profile, extra=extra)
            else:
                send_telegram("請從鍵盤揀選 MBTI，或者打「跳過」")
            return

    # Practice / Review session
    session = load_session()
    state   = session.get("state", "") if session else ""

    if state == "waiting_response" and not text.startswith("/"):
        handle_user_response(text)
        return

    if state == "waiting_review" and not text.startswith("/"):
        clear_session()
        profile = load_profile() or {}
        send_telegram("🔍 AI 分析緊你嘅面試對話，稍等⋯⋯")
        result = analyze_conversation(text, profile)
        send_telegram(result, reply_markup={"inline_keyboard": [[
            {"text": "🎯 繼續練習", "callback_data": "practice_new"},
        ]]})
        return

    if state == "negotiate_start" and not text.startswith("/"):
        session["offer_details"] = text.strip()
        session["round_num"]     = 0
        session["history"]       = []
        session["state"]         = "negotiate_session"
        save_session(session)
        send_telegram(
            "✅ Offer 已記錄！而家你可以開始同 HR 傾，打你想講嘅第一句：",
            reply_markup={"inline_keyboard": [[{"text": "🏁 結束談判", "callback_data": "negotiate_end"}]]}
        )
        return

    if state == "negotiate_session":
        if text.strip() == "結束" or cmd(text, "/negotiate"):
            send_telegram("📊 生成談判總結⋯⋯")
            summary = generate_negotiate_summary(session.get("history", []))
            send_telegram(summary, reply_markup={"inline_keyboard": [[
                {"text": "🎯 繼續練習", "callback_data": "practice_new"},
            ]]})
            _record_negotiate_summary(session, summary)
            clear_session()
            return

        round_num = session.get("round_num", 0) + 1
        history   = session.get("history", [])
        reply = generate_negotiate_response(session.get("offer_details", ""), text.strip(), round_num, history)
        history.append({"user": text.strip(), "hr": extract_negotiate_reply(reply)})
        session["round_num"] = round_num
        session["history"]   = history
        save_session(session)
        send_telegram(reply, reply_markup={"inline_keyboard": [[
            {"text": "🏁 結束談判", "callback_data": "negotiate_end"},
        ]]})
        return

    if state == "debrief_input" and not text.startswith("/"):
        job_info = session.get("job_info")
        clear_session()
        send_telegram("🎙️ AI 分析緊你嘅面試表現⋯⋯")
        result = generate_debrief(job_info, text.strip())
        _record_debrief_result(job_info, result)
        if job_info:
            send_telegram(result)
            handle_update_status_menu(job_info["id"])
        else:
            send_telegram(result, reply_markup={"inline_keyboard": [[
                {"text": "🎯 繼續練習", "callback_data": "practice_new"},
            ]]})
        return

    # Commands
    if cmd(text, "/start") or cmd(text, "/help"):
        profile = load_profile()
        if not profile or (not profile.get("job_title") and not profile.get("industry")):
            send_telegram(
                "👋 歡迎！AI 面試教練幫你針對練習，提升面試表現。\n\n"
                "首先設定你嘅背景，練習會更貼近你嘅實際情況：",
                reply_markup={"inline_keyboard": [
                    [{"text": "📄 上傳 CV（PDF / DOCX）", "callback_data": "onboard_cv"}],
                    [{"text": "✏️ 自己輸入背景",          "callback_data": "onboard_manual"}],
                ]}
            )
            return
        send_telegram(
            "🎓 AI 面試教練\n\n"
            "📋 *求職追蹤*\n"
            "/addjob — 新增求職申請\n"
            "/listjobs — 睇所有申請 + AI 面試問題\n\n"
            "🎯 *練習*\n"
            "/practice — 隨機面試練習\n"
            "/practice 初級／中級／高級 — 指定難度\n"
            "/drill — 針對特定題型練習\n"
            "/review — 貼真實面試答案，AI 分析\n"
            "/negotiate — 薪酬談判 role-play\n"
            "/debrief — 面試後覆盤分析\n\n"
            "📊 *統計*\n"
            "/stats — 我的進度\n"
            "/streak — 練習連續天數\n"
            "/tip — 今日面試技巧\n\n"
            "⚙️ *設定*\n"
            "/setup — 更改設定\n"
            "/mystatus — 我的設定\n\n"
            "📄 上傳 PDF / DOCX resume — AI 自動分析並個人化練習"
        )
        return

    if cmd(text, "/setup"):
        clear_setup_session()
        send_telegram(
            "⚙️ 更新背景設定：",
            reply_markup={"inline_keyboard": [
                [{"text": "📄 重新上傳 CV", "callback_data": "onboard_cv"}],
                [{"text": "✏️ 手動輸入",    "callback_data": "onboard_manual"}],
            ]}
        )
        return

    if cmd(text, "/practice"):
        profile = load_profile()
        if not profile or (not profile.get("job_title") and not profile.get("industry")):
            send_industry_keyboard(intro=True)
            return
        parts = text.split(maxsplit=1)
        extra = parts[1].strip() if len(parts) > 1 else None
        diff  = extra if extra in DIFFICULTY_LEVELS else None
        send_telegram("🎯 生成面試場景⋯⋯")
        start_practice(difficulty=diff)
        return

    if cmd(text, "/drill"):
        handle_drill_menu()
        return

    if cmd(text, "/review"):
        save_session({"state": "waiting_review"})
        send_telegram(
            "📋 貼上你嘅真實面試問答記錄，AI 幫你分析失分點：\n\n"
            "（格式：面試官問：xxx\n我答：xxx）"
        )
        return

    if cmd(text, "/negotiate"):
        save_session({"state": "negotiate_start"})
        send_telegram(
            "🤝 薪酬談判練習\n\n貼你收到嘅 offer details（職位、公司、提供月薪、其他 package）：",
            reply_markup={"inline_keyboard": [[{"text": "❌ 取消", "callback_data": "negotiate_cancel"}]]}
        )
        return

    if cmd(text, "/debrief"):
        jobs = load_jobs()
        if jobs:
            kb = [
                [{"text": f"{STATUS_EMOJI.get(j.get('status','Applied'),'')} {j['company']} — {j['role']}",
                  "callback_data": f"debrief_job_{j['id']}"}]
                for j in jobs
            ]
            kb.append([{"text": "⏭️ 跳過，唔連結特定工", "callback_data": "debrief_job_skip"}])
            save_session({"state": "debrief_job_select"})
            send_telegram("🎙️ 係邊份工嘅面試？", reply_markup={"inline_keyboard": kb})
        else:
            save_session({"state": "debrief_input", "job_info": None})
            send_telegram(DEBRIEF_PROMPT)
        return

    if cmd(text, "/mbti"):
        cmd_mbti(text)
        return

    if cmd(text, "/stats"):
        handle_stats()
        return

    if cmd(text, "/streak"):
        data   = load_stats()
        streak = data.get("streak", {})
        count  = streak.get("count", 0)
        total  = data.get("total_sessions", 0)
        emoji  = "🔥" if count >= 7 else "💪" if count >= 3 else "🌱"
        send_telegram(
            f"{emoji} 連續練習：{count} 日\n"
            f"總練習次數：{total} 次\n\n"
            f"每日練習，面試信心係咁升！"
        )
        return

    if cmd(text, "/tip"):
        send_telegram(f"💡 今日面試技巧\n\n{get_daily_tip()}")
        return

    if cmd(text, "/mystatus"):
        profile = load_profile() or {}
        data    = load_stats()
        total   = data.get("total_sessions", 0)
        streak  = data.get("streak", {}).get("count", 0)
        mbti    = profile.get("mbti", "未設定")
        note    = ""
        if profile.get("mbti") and profile["mbti"] in MBTI_COACHING:
            note = f"\n💡 面試盲點：{MBTI_COACHING[profile['mbti']]['watch_out']}"
        send_telegram(
            f"⚙️ 我的設定\n\n"
            f"🎯 目標職位：{profile.get('job_title','未設定')}\n"
            f"🏭 行業：{profile.get('industry','未設定')}\n"
            f"🧠 MBTI：{mbti}{note}\n\n"
            f"📊 總練習：{total} 次  🔥 連續：{streak} 日\n\n"
            f"更改設定：/setup",
            reply_markup={"inline_keyboard": [[{"text": "🎯 開始練習", "callback_data": "practice_new"}]]}
        )
        return

    if cmd(text, "/addjob"):
        handle_addjob_start()
        return

    if cmd(text, "/listjobs") or cmd(text, "/jobs"):
        handle_listjobs()
        return

    # 未知輸入：引導
    profile = load_profile()
    if not profile or (not profile.get("job_title") and not profile.get("industry")):
        send_telegram(
            "先設定你嘅背景，練習會更個人化 👇",
            reply_markup={"inline_keyboard": [
                [{"text": "📄 上傳 CV", "callback_data": "onboard_cv"}],
                [{"text": "✏️ 手動輸入", "callback_data": "onboard_manual"}],
            ]}
        )
    else:
        send_telegram(
            "輸入 /practice 開始練習，/help 睇所有指令 😊",
            reply_markup={"inline_keyboard": [[{"text": "🎯 開始練習", "callback_data": "practice_new"}]]}
        )


# ── Routes ────────────────────────────────────────────────────────

@app.route("/api/webhook", methods=["POST"])
def webhook():
    update = request.json or {}

    # 抽取 chat_id，令 send_telegram 知道發去邊
    trigger = ""
    if "callback_query" in update:
        chat_id = update["callback_query"]["message"]["chat"]["id"]
        set_current_chat_id(chat_id)
        trigger = update["callback_query"].get("data", "")
    elif "message" in update:
        chat_id = update["message"]["chat"]["id"]
        set_current_chat_id(chat_id)
        trigger = update["message"].get("text") or "[document]"

    try:
        if "callback_query" in update:
            handle_callback(update["callback_query"])
        elif "message" in update:
            document = update["message"].get("document")
            text     = update["message"].get("text", "").strip()
            if document:
                handle_document(document)
            elif text:
                handle_message(text)
    except Exception as e:
        tb = traceback.format_exc()
        print(tb)
        try:
            send_telegram(f"⚠️ Bot error\n指令：{trigger}\n錯誤：{type(e).__name__}: {e}")
        except Exception:
            pass

    return jsonify({"ok": True})


@app.route("/api/set_webhook", methods=["GET"])
def set_webhook():
    """部署後訪問呢個 URL — 自動設定 webhook + 更新指令清單。"""
    host = request.host_url.rstrip("/")
    url  = f"{host}/api/webhook"

    # 1. 設定 webhook
    wh_resp = req.post(
        f"https://api.telegram.org/bot{TOKEN()}/setWebhook",
        json={"url": url, "allowed_updates": ["message", "callback_query"]},
        timeout=10,
    ).json()

    # 2. 更新 Telegram 指令清單
    commands = [
        {"command": "addjob",   "description": "新增求職申請記錄"},
        {"command": "listjobs", "description": "查看所有申請 + AI 面試問題"},
        {"command": "practice", "description": "隨機面試練習"},
        {"command": "drill",    "description": "針對特定題型練習"},
        {"command": "stats",    "description": "我的進度報告"},
        {"command": "streak",   "description": "練習連續天數"},
        {"command": "tip",      "description": "今日面試技巧"},
        {"command": "review",   "description": "貼真實面試答案，AI 分析"},
        {"command": "negotiate", "description": "薪酬談判 role-play"},
        {"command": "debrief",  "description": "面試後覆盤分析"},
        {"command": "mbti",     "description": "做 MBTI 檢測 / 直接輸入 MBTI"},
        {"command": "setup",    "description": "設定目標職位 + MBTI"},
        {"command": "mystatus", "description": "查看我的設定"},
        {"command": "help",     "description": "指令說明（提示：可直接上傳 PDF resume）"},
    ]
    cmd_resp = req.post(
        f"https://api.telegram.org/bot{TOKEN()}/setMyCommands",
        json={"commands": commands},
        timeout=10,
    ).json()

    return jsonify({
        "webhook_url": url,
        "webhook": wh_resp,
        "commands": cmd_resp,
    })


@app.route("/", methods=["GET"])
def health():
    return "AI Interview Coach Bot ✅", 200


if __name__ == "__main__":
    app.run(debug=False)
