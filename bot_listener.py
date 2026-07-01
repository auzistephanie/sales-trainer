"""AI 面試教練 Bot — Telegram 對話管理、指令路由、進度追蹤。"""

import logging
import logging.handlers
import os
import sys
import re
import time
import threading
import requests
from pathlib import Path
from datetime import datetime, timedelta
from dotenv import load_dotenv

BASE_DIR = Path(__file__).parent
os.chdir(BASE_DIR)
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

load_dotenv(BASE_DIR / ".env")

_log_file = BASE_DIR / "bot.log"
_handler  = logging.handlers.RotatingFileHandler(
    _log_file, maxBytes=5 * 1024 * 1024, backupCount=3, encoding="utf-8"
)
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(message)s",
    datefmt="%H:%M:%S",
    handlers=[_handler, logging.StreamHandler(sys.stdout)],
)
log = logging.getLogger(__name__)

from interview_trainer import (
    generate_scenario, evaluate_response, analyze_conversation,
    QUESTION_TYPES, INDUSTRIES, DIFFICULTY_LEVELS, MBTI_COACHING,
    get_daily_tip, generate_job_questions, generate_job_tips,
    parse_resume, calculate_cv_health, format_cv_health_message,
    generate_salary_benchmark, parse_salary_input,
    extract_job_from_url,
    generate_cover_letter_from_jd, generate_tailored_cv_content, build_cv_docx,
    generate_negotiate_response, generate_negotiate_summary, extract_negotiate_reply,
    generate_debrief,
)
from utils import (
    load_stats, save_stats,
    load_session, save_session, clear_session,
    load_profile, save_profile,
    load_setup_session, save_setup_session, clear_setup_session,
    load_jobs, save_jobs,
    load_addjob_session, save_addjob_session, clear_addjob_session,
    load_cv_text, save_cv_text,
    send_telegram, send_document,
)

# ── 免費 session 上限 ──────────────────────────────────────────────
FREE_SESSION_LIMIT = 5
UPGRADE_MSG = (
    "🎓 你已完成 {n} 次免費練習！\n\n"
    "升級 Premium 解鎖無限練習 + 詳細進度分析：\n"
    "👉 [加入等候名單](https://t.me/hkinterviewbot)（暫定 $68/月）\n\n"
    "或者繼續試用，每日送 1 次額外練習 🎁"
)


# ── 基礎工具 ──────────────────────────────────────────────────────

def answer_callback(cb_id: str, text: str = ""):
    token = os.getenv("TELEGRAM_BOT_TOKEN")
    requests.post(
        f"https://api.telegram.org/bot{token}/answerCallbackQuery",
        json={"callback_query_id": cb_id, "text": text, "show_alert": False},
        timeout=10,
    )


JOB_STATUSES = ["Applied", "Phone Screen", "1st Interview", "2nd Interview", "Offer", "Rejected"]
STATUS_EMOJI  = {"Applied": "📝", "Phone Screen": "📞", "1st Interview": "🤝",
                 "2nd Interview": "🔁", "Offer": "🎉", "Rejected": "❌"}

DEBRIEF_PROMPT = (
    "請描述面試過程。包括：\n"
    "- 問咗咩問題\n"
    "- 你點答\n"
    "- Interviewer 嘅反應\n"
    "- 你覺得咩位答得唔好"
)


def register_commands():
    token = os.getenv("TELEGRAM_BOT_TOKEN")
    commands = [
        {"command": "practice",   "description": "隨機面試練習"},
        {"command": "drill",      "description": "針對特定題型練習"},
        {"command": "addjob",     "description": "新增求職申請記錄"},
        {"command": "listjobs",   "description": "查看所有申請 + 狀態"},
        {"command": "stats",      "description": "我的進度報告"},
        {"command": "streak",     "description": "練習連續天數"},
        {"command": "tip",        "description": "今日面試技巧"},
        {"command": "review",     "description": "貼真實面試答案，AI 分析失分點"},
        {"command": "negotiate",  "description": "薪酬談判 role-play"},
        {"command": "debrief",    "description": "面試後覆盤分析"},
        {"command": "setup",      "description": "設定我的目標職位 + MBTI"},
        {"command": "mystatus",   "description": "查看我的設定"},
        {"command": "help",       "description": "指令說明"},
    ]
    requests.post(
        f"https://api.telegram.org/bot{token}/setMyCommands",
        json={"commands": commands},
        timeout=10,
    )


# ── 統計記錄 ──────────────────────────────────────────────────────

def record_score(qtype_name: str, score: int):
    data    = load_stats()
    scores  = data.setdefault("qtype_scores", {})
    history = scores.setdefault(qtype_name, [])
    history.append(score)
    scores[qtype_name] = history[-20:]
    data["total_sessions"] = data.get("total_sessions", 0) + 1

    today     = datetime.now().strftime("%Y-%m-%d")
    yesterday = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")
    streak    = data.setdefault("streak", {"last_date": "", "count": 0})
    if streak["last_date"] == today:
        pass
    elif streak["last_date"] == yesterday:
        streak["count"] += 1
        streak["last_date"] = today
    else:
        streak["count"]    = 1
        streak["last_date"] = today

    save_stats(data)
    avg = sum(history) / len(history)
    return avg, len(history)


def check_free_limit() -> bool:
    """True = 可以繼續，False = 超出免費次數。"""
    data  = load_stats()
    total = data.get("total_sessions", 0)
    # 每日免費 bonus：今日已免費練習 < 1
    today_bonus = data.get("today_bonus_date") == datetime.now().strftime("%Y-%m-%d")
    if total < FREE_SESSION_LIMIT or today_bonus:
        return True
    return False


def use_daily_bonus():
    """記錄今日已用免費 bonus。"""
    data = load_stats()
    data["today_bonus_date"] = datetime.now().strftime("%Y-%m-%d")
    save_stats(data)


# ── Stats 報告 ────────────────────────────────────────────────────

def handle_stats():
    data   = load_stats()
    scores = data.get("qtype_scores", {})
    total  = data.get("total_sessions", 0)
    streak = data.get("streak", {})

    if not scores:
        send_telegram("未有練習記錄，用 /practice 開始第一次練習！")
        return

    ranked = sorted(
        [(name, sum(sc)/len(sc), len(sc)) for name, sc in scores.items() if sc],
        key=lambda x: x[1], reverse=True
    )

    profile = load_profile()
    mbti_line = f"🧠 MBTI：{profile.get('mbti', '未設定')}  " if profile.get("mbti") else ""
    job_line  = f"🎯 目標：{profile.get('job_title', '未設定')}" if profile.get("job_title") else ""

    lines = [
        "📊 我的面試訓練進度",
        f"{mbti_line}{job_line}",
        f"總練習：{total} 次  |  連續 {streak.get('count', 0)} 日\n",
        "【各題型掌握度】",
    ]
    for name, avg, count in ranked:
        filled = "█" * round(avg)
        empty  = "░" * (4 - round(avg))
        pct    = round(avg / 4 * 100)
        lines.append(f"{filled}{empty}  {name}  {pct}%（{count}次）")

    weak   = [x for x in ranked if x[1] < 2.5]
    strong = [x for x in ranked if x[1] >= 3.5]

    if weak:
        lines.append(f"\n⚠️ 薄弱項目：{', '.join(x[0] for x in weak[:3])}")
        lines.append("用 /drill 針對練習 👆")
    if strong:
        lines.append(f"\n💪 已掌握：{', '.join(x[0] for x in strong[:3])}")

    send_telegram("\n".join(lines))


# ── 今日技巧 ──────────────────────────────────────────────────────

def handle_tip():
    send_telegram(f"💡 今日面試技巧\n\n{get_daily_tip()}")


# ── Profile Setup ─────────────────────────────────────────────────

# 行業選單（14 種，每行 2 個）
SETUP_INDUSTRY_LIST = [
    ("💰", "金融/投行",    "金融／投資銀行"),
    ("💻", "科技/IT",      "科技／軟件開發"),
    ("📣", "市場/廣告",    "市場營銷／廣告"),
    ("🧩", "管理諮詢",     "管理諮詢"),
    ("🛍️", "零售/酒店",   "零售／酒店管理"),
    ("🚀", "初創",         "初創公司"),
    ("🏥", "醫療/健康科技", "醫療／藥劑／健康科技"),
    ("⚖️", "法律/合規",   "法律／合規"),
    ("👥", "人力資源",     "人力資源／招聘"),
    ("📚", "教育/培訓",    "教育／培訓"),
    ("🚚", "物流/供應鏈",  "物流／供應鏈"),
    ("🏗️", "地產/建築",  "地產／建築工程"),
    ("🎬", "傳媒/創意",   "傳媒／娛樂／創意"),
    ("🏛️", "政府/NGO",   "政府／NGO／公共服務"),
]

# MBTI 選單（16 種，4×4）
MBTI_LIST = [
    "INTJ", "INTP", "ENTJ", "ENTP",
    "INFJ", "INFP", "ENFJ", "ENFP",
    "ISTJ", "ISFJ", "ESTJ", "ESFJ",
    "ISTP", "ISFP", "ESTP", "ESFP",
]


def handle_setup_start(intro: bool = True):
    """Step 1：揀行業。"""
    keyboard = []
    row = []
    for i, (icon, short, _) in enumerate(SETUP_INDUSTRY_LIST):
        row.append({"text": f"{icon} {short}", "callback_data": f"setup_ind_{i}"})
        if len(row) == 2:
            keyboard.append(row)
            row = []
    if row:
        keyboard.append(row)
    keyboard.append([{"text": "✏️ 自定行業", "callback_data": "setup_ind_custom"}])

    msg = (
        "👋 歡迎使用 AI 面試教練！\n先設定你嘅背景，令練習更貼近你嘅情況。\n\n🏭 目標行業係？"
        if intro else
        "🏭 揀你嘅目標行業："
    )
    send_telegram(msg, reply_markup={"inline_keyboard": keyboard})


def _handle_negotiate_turn(session: dict, text: str):
    """背景 thread：生成 HR 回應 + 評分，繼續談判。"""
    round_num = session.get("round_num", 0) + 1
    history   = session.get("history", [])
    reply = generate_negotiate_response(session.get("offer_details", ""), text, round_num, history)
    history.append({"user": text, "hr": extract_negotiate_reply(reply)})
    session["round_num"] = round_num
    session["history"]   = history
    save_session(session)
    send_telegram(reply, reply_markup={"inline_keyboard": [[
        {"text": "🏁 結束談判", "callback_data": "negotiate_end"},
    ]]})


def _handle_negotiate_summary(history: list):
    """背景 thread：生成談判總結。"""
    send_telegram("📊 生成談判總結⋯⋯")
    summary = generate_negotiate_summary(history)
    send_telegram(summary, reply_markup={"inline_keyboard": [[
        {"text": "🎯 繼續練習", "callback_data": "practice_new"},
    ]]})


def _handle_debrief_result(job_info, text: str):
    """背景 thread：生成面試覆盤分析，如有連結職位就跟住問更新狀態。"""
    send_telegram("🎙️ AI 分析緊你嘅面試表現⋯⋯")
    result = generate_debrief(job_info, text)
    if job_info:
        send_telegram(result)
        handle_update_status_menu(job_info["id"])
    else:
        send_telegram(result, reply_markup={"inline_keyboard": [[
            {"text": "🎯 繼續練習", "callback_data": "practice_new"},
        ]]})


def _finish_salary_step(profile: dict, expected_salary: str):
    """背景 thread：生成薪酬 benchmark，然後進入 MBTI 步。"""
    benchmark = generate_salary_benchmark(
        profile.get("job_title", "未指定職位"), expected_salary, profile.get("industry", "")
    )
    send_telegram(f"💰 HK 市場薪酬參考（2026）\n\n{benchmark}")
    save_setup_session({"state": "setup_mbti"})
    send_mbti_keyboard()


def send_mbti_keyboard():
    """顯示 MBTI 選擇鍵盤（4×4）。"""
    keyboard = []
    for i in range(0, 16, 4):
        row = [{"text": t, "callback_data": f"setup_mbti_{t}"} for t in MBTI_LIST[i:i+4]]
        keyboard.append(row)
    keyboard.append([{"text": "⏭️ 唔知 / 跳過", "callback_data": "setup_mbti_skip"}])
    send_telegram(
        "🧠 你嘅 MBTI 係？（唔知可以跳過，之後 /setup 再改）\n"
        "唔識做 MBTI 測試？👉 https://www.16personalities.com/ch",
        reply_markup={"inline_keyboard": keyboard},
    )


def handle_document(document: dict):
    """處理用戶上傳嘅 resume（PDF / DOCX）。"""
    mime  = document.get("mime_type", "")
    fname = document.get("file_name", "").lower()

    is_pdf  = mime == "application/pdf" or fname.endswith(".pdf")
    is_docx = mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" \
              or fname.endswith(".docx")

    if not is_pdf and not is_docx:
        send_telegram("⚠️ 請上傳 PDF 或 Word (.docx) 格式嘅 resume。")
        return

    send_telegram("📄 收到 resume！AI 分析緊，稍等⋯⋯")

    token = os.getenv("TELEGRAM_BOT_TOKEN")
    file_info = requests.get(
        f"https://api.telegram.org/bot{token}/getFile",
        params={"file_id": document["file_id"]}, timeout=10,
    ).json()
    if not file_info.get("ok"):
        send_telegram("❌ 下載失敗，請重試。")
        return

    file_url   = f"https://api.telegram.org/file/bot{token}/{file_info['result']['file_path']}"
    file_bytes = requests.get(file_url, timeout=30).content

    resume_text = ""
    try:
        if is_pdf:
            import pdfplumber, io
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                resume_text = "\n".join(p.extract_text() or "" for p in pdf.pages)
        else:
            import docx, io
            doc = docx.Document(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = "  |  ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        parts.append(row_text)
            resume_text = "\n".join(parts)
    except Exception as e:
        send_telegram(f"❌ 解析文件失敗：{e}\n請確保唔係掃描版圖片 PDF。")
        return

    if not resume_text.strip():
        send_telegram("⚠️ 未能提取文字，可能係掃描版 PDF。請試用文字版 PDF 或 .docx。")
        return

    parsed = parse_resume(resume_text)
    if not parsed:
        send_telegram("❌ AI 分析失敗，請重試。")
        return

    save_cv_text(resume_text)
    profile = load_profile()
    if parsed.get("job_title"):       profile["job_title"] = parsed["job_title"]
    if parsed.get("industry"):        profile["industry"]  = parsed["industry"]
    if parsed.get("exp_years"):       profile["exp_years"] = str(parsed["exp_years"])
    if parsed.get("current_company"): profile["company"]   = parsed["current_company"]
    if parsed.get("key_skills"):      profile["key_skills"] = parsed["key_skills"]
    if parsed.get("education"):       profile["education"]  = parsed["education"]

    health = calculate_cv_health(resume_text)
    profile["cv_health_score"] = health["total"]
    save_profile(profile)

    lines = [
        "✅ *Resume 分析完成！*\n",
        f"📝 {parsed.get('summary', '')}",
        "",
        f"🎯 職位：{parsed.get('job_title', '未識別')}",
        f"🏭 行業：{parsed.get('industry', '未識別')}",
        f"⏱️ 年資：{parsed.get('exp_years', '?')} 年",
    ]
    if parsed.get("current_company"): lines.append(f"🏢 公司：{parsed['current_company']}")
    if parsed.get("key_skills"):      lines.append(f"🛠️ 技能：{parsed['key_skills']}")
    if parsed.get("education"):       lines.append(f"🎓 學歷：{parsed['education']}")

    setup = load_setup_session()
    onboarding = bool(setup) and setup.get("state") == "setup_cv_upload"

    if not onboarding:
        lines.append("\n面試練習會根據你嘅背景個人化！")
    send_telegram("\n".join(lines))
    send_telegram(format_cv_health_message(health))

    if onboarding:
        save_setup_session({"state": "setup_salary"})
        send_telegram("💰 仲有最後一步！你目標月薪期望大概係幾多？（例如：38000 或 38K）")
    else:
        send_telegram(
            "繼續做咩？",
            reply_markup={"inline_keyboard": [[
                {"text": "🎯 立即練習", "callback_data": "practice_new"},
                {"text": "⚙️ 修改設定",  "callback_data": "do_setup"},
            ]]}
        )


# ── Drill 選單 ────────────────────────────────────────────────────

def handle_drill_menu():
    keyboard = []
    row = []
    for i, q in enumerate(QUESTION_TYPES):
        row.append({"text": q["name"], "callback_data": f"drill_{q['name']}"})
        if len(row) == 2:
            keyboard.append(row)
            row = []
    if row:
        keyboard.append(row)
    keyboard.append([{"text": "🎲 隨機場景", "callback_data": "practice_new"}])
    send_telegram("🎯 選擇你想針對練習的題型：", reply_markup={"inline_keyboard": keyboard})


# ── 練習 Session ──────────────────────────────────────────────────

def start_practice(force_qtype: str = None, force_industry: str = None, difficulty: str = None):
    if not check_free_limit():
        data  = load_stats()
        total = data.get("total_sessions", 0)
        kb = {"inline_keyboard": [[
            {"text": "🎁 領取今日免費練習", "callback_data": "claim_bonus"},
            {"text": "📊 睇進度",           "callback_data": "show_stats"},
        ]]}
        send_telegram(UPGRADE_MSG.format(n=total), reply_markup=kb)
        return

    if not force_industry:
        profile = load_profile()
        force_industry = profile.get("industry")

    display, scenario = generate_scenario(force_qtype, force_industry, difficulty)
    save_session({"state": "waiting_response", "scenario": scenario})
    send_telegram(display)


def handle_user_response(user_text: str):
    session = load_session()
    if not session or session.get("state") != "waiting_response":
        send_telegram("唔係練習模式。用 /practice 開始新練習！")
        return

    clear_session()
    scenario  = session["scenario"]
    qtype_name = scenario["qtype"]["name"]

    send_telegram("🤔 AI 評估緊你嘅回答，稍等⋯⋯")
    feedback = evaluate_response(user_text, scenario, profile=load_profile())

    # 解析分數
    match = re.search(r"評分[：:]\s*([1-4])", feedback)
    score = int(match.group(1)) if match else None
    if score:
        record_score(qtype_name, score)

    replay_kb = {"inline_keyboard": [
        [
            {"text": "🔄 再練一個",              "callback_data": "practice_new"},
            {"text": f"🎯 再練「{qtype_name}」", "callback_data": f"drill_{qtype_name}"},
        ],
        [
            {"text": "📊 睇進度",               "callback_data": "show_stats"},
            {"text": "💡 今日技巧",             "callback_data": "show_tip"},
        ],
    ]}
    send_telegram(feedback, reply_markup=replay_kb)


# ── Job Application Tracker ───────────────────────────────────────

def _auto_add_job_from_url(url: str):
    """後台：fetch URL → 抽取 company/role/jd → save job + 生成 CV + Cover Letter。"""
    from datetime import date as _date

    # 1. 抽取職位資料
    info    = extract_job_from_url(url)
    company = info.get("company", "").strip()
    role    = info.get("role", "").strip()
    jd      = info.get("jd", "").strip()

    if not company or not role:
        send_telegram(
            "⚠️ 未能自動讀取職位資料（可能需要登入或頁面受保護）。\n\n"
            "用 /addjob 手動新增，貼上公司、職位同 JD。"
        )
        return

    # 2. Save job to tracker
    jobs   = load_jobs()
    new_id = (max(j["id"] for j in jobs) + 1) if jobs else 1
    jobs.append({
        "id":           new_id,
        "company":      company,
        "role":         role,
        "jd":           jd,
        "link":         url,
        "applied_date": str(_date.today()),
        "status":       "Applied",
    })
    save_jobs(jobs)
    send_telegram(
        f"✅ 已新增：{company} — {role}\n"
        f"⏳ 正在生成 Tailored CV 同 Cover Letter..."
    )

    # 3. 讀 CV — 冇 CV 就提示上傳
    cv_text = load_cv_text()
    if not cv_text:
        send_telegram(
            "⚠️ 未有 CV 記錄，無法生成 Tailored CV。\n\n"
            "請發你的 CV（PDF 或 .docx）比我，我會自動解析。"
        )
        return

    # 4. 生成 Cover Letter
    cover = generate_cover_letter_from_jd(cv_text, jd, company, role)
    send_telegram(f"📝 *Cover Letter — {company}*\n\n{cover}")

    # 5. 生成 Tailored CV .docx
    cv_data = generate_tailored_cv_content(cv_text, jd, company, role)
    if cv_data:
        docx_bytes = build_cv_docx(cv_data, company, role)
        filename   = f"CV_{company.replace(' ', '_')}_{role.replace(' ', '_')}.docx"
        send_document(docx_bytes, filename, caption=f"📄 Tailored CV — {role} @ {company}")


def handle_addjob_start():
    clear_addjob_session()
    save_addjob_session({"state": "addjob_company"})
    send_telegram("📋 新增求職申請\n\n🏢 公司名稱係？（例如：TransUnion）")


def handle_listjobs():
    jobs = load_jobs()
    if not jobs:
        send_telegram(
            "未有申請記錄。用 /addjob 新增第一個！",
            reply_markup={"inline_keyboard": [[{"text": "➕ 新增申請", "callback_data": "addjob_start"}]]}
        )
        return

    lines = ["📋 我的求職申請\n"]
    keyboard = []
    for j in jobs:
        jid    = j["id"]
        emoji  = STATUS_EMOJI.get(j.get("status", "Applied"), "📝")
        date_s = j.get("applied_date", "")
        lines.append(f"{jid}. {emoji} *{j['company']}* — {j['role']}")
        lines.append(f"   狀態：{j.get('status', 'Applied')}  |  申請日：{date_s}\n")
        keyboard.append([
            {"text": f"#{jid} Questions", "callback_data": f"job_q_{jid}"},
            {"text": f"#{jid} Tips",      "callback_data": f"job_tips_{jid}"},
        ])
        keyboard.append([
            {"text": f"#{jid} Practice",        "callback_data": f"job_practice_{jid}"},
            {"text": f"#{jid} Update Status",   "callback_data": f"job_updatestatus_{jid}"},
        ])
        keyboard.append([
            {"text": f"#{jid} 🤝 Negotiate", "callback_data": f"job_negotiate_{jid}"},
        ])

    keyboard.append([{"text": "➕ 新增申請", "callback_data": "addjob_start"}])
    send_telegram("\n".join(lines), reply_markup={"inline_keyboard": keyboard})


def handle_job_questions(job_id: int):
    jobs = load_jobs()
    job  = next((j for j in jobs if j["id"] == job_id), None)
    if not job:
        send_telegram(f"找不到申請 #{job_id}，用 /listjobs 查看。")
        return
    send_telegram(f"🤔 AI 生成 #{job_id} {job['company']} — {job['role']} 面試問題⋯⋯")
    result = generate_job_questions(job)
    send_telegram(
        f"❓ *面試問題 — {job['company']} / {job['role']}*\n\n{result}",
        reply_markup={"inline_keyboard": [[
            {"text": "💼 Key Tips",      "callback_data": f"job_tips_{job_id}"},
            {"text": "🎯 開始練習",      "callback_data": f"job_practice_{job_id}"},
        ]]}
    )


def handle_job_tips(job_id: int):
    jobs = load_jobs()
    job  = next((j for j in jobs if j["id"] == job_id), None)
    if not job:
        send_telegram(f"找不到申請 #{job_id}。")
        return
    send_telegram(f"🤔 AI 生成 #{job_id} {job['company']} Key Talking Points⋯⋯")
    result = generate_job_tips(job)
    send_telegram(
        f"💼 *Key Talking Points — {job['company']} / {job['role']}*\n\n{result}",
        reply_markup={"inline_keyboard": [[
            {"text": "❓ 面試問題",  "callback_data": f"job_q_{job_id}"},
            {"text": "🎯 開始練習", "callback_data": f"job_practice_{job_id}"},
        ]]}
    )


def handle_update_status_menu(job_id: int):
    jobs = load_jobs()
    job  = next((j for j in jobs if j["id"] == job_id), None)
    if not job:
        send_telegram(f"找不到申請 #{job_id}。")
        return
    keyboard = []
    row = []
    for s in JOB_STATUSES:
        row.append({"text": f"{STATUS_EMOJI[s]} {s}", "callback_data": f"job_status_{job_id}_{s}"})
        if len(row) == 2:
            keyboard.append(row); row = []
    if row: keyboard.append(row)
    send_telegram(
        f"📌 更新 *{job['company']} — {job['role']}* 狀態：",
        reply_markup={"inline_keyboard": keyboard}
    )


# ── Callback ──────────────────────────────────────────────────────

def handle_callback(cb: dict):
    data = cb.get("data", "")

    if data == "practice_new":
        answer_callback(cb["id"], "生成新場景⋯⋯")
        threading.Thread(target=start_practice, daemon=True).start()

    elif data == "claim_bonus":
        answer_callback(cb["id"], "🎁 今日免費練習已解鎖！")
        use_daily_bonus()
        threading.Thread(target=start_practice, daemon=True).start()

    elif data == "show_stats":
        answer_callback(cb["id"])
        handle_stats()

    elif data == "show_tip":
        answer_callback(cb["id"])
        handle_tip()

    elif data == "do_setup":
        answer_callback(cb["id"])
        clear_setup_session()
        handle_setup_start(intro=False)

    elif data == "onboard_cv":
        answer_callback(cb["id"])
        save_setup_session({"state": "setup_cv_upload"})
        send_telegram(
            "📄 請直接將你嘅 CV 拖入呢個 chat（PDF 或 Word .docx）。\n\n"
            "AI 會自動分析你嘅背景，然後開始個人化練習！"
        )

    elif data == "onboard_manual":
        answer_callback(cb["id"])
        clear_setup_session()
        handle_setup_start(intro=False)

    # ── Setup：行業 ──
    elif data == "setup_ind_custom":
        answer_callback(cb["id"])
        save_setup_session({"state": "setup_industry_custom"})
        send_telegram("✏️ 打出你嘅目標行業（例如：航空、網絡安全、遊戲）：")

    elif data.startswith("setup_ind_"):
        idx = int(data[len("setup_ind_"):])
        if idx < len(SETUP_INDUSTRY_LIST):
            _, _, full_name = SETUP_INDUSTRY_LIST[idx]
            profile = load_profile()
            profile["industry"] = full_name
            save_profile(profile)
            answer_callback(cb["id"], f"✅ {full_name}")
            save_setup_session({"state": "setup_jobtitle"})
            send_telegram(f"✅ 行業：{full_name}\n\n🎯 你嘅目標職位係？（直接打，例如：Product Manager）")
        else:
            answer_callback(cb["id"], "出錯，請再試")

    # ── Setup：MBTI ──
    elif data.startswith("setup_mbti_"):
        mbti_val = data[len("setup_mbti_"):]
        profile  = load_profile()
        if mbti_val == "skip":
            answer_callback(cb["id"], "已跳過")
            clear_setup_session()
            _send_setup_done(profile)
        else:
            profile["mbti"] = mbti_val
            save_profile(profile)
            answer_callback(cb["id"], f"✅ {mbti_val}")
            clear_setup_session()
            coaching = MBTI_COACHING.get(mbti_val, {})
            note = f"\n\n💡 {mbti_val} 面試特點：{coaching.get('watch_out', '')}" if coaching else ""
            _send_setup_done(profile, extra=note)

    # ── Drill ──
    elif data.startswith("drill_"):
        qtype_name = data[6:]
        answer_callback(cb["id"], f"生成「{qtype_name}」場景⋯⋯")
        threading.Thread(
            target=start_practice,
            kwargs={"force_qtype": qtype_name},
            daemon=True,
        ).start()

    # ── Job Tracker ──
    elif data == "addjob_start":
        answer_callback(cb["id"])
        handle_addjob_start()

    elif data.startswith("job_q_"):
        job_id = int(data[len("job_q_"):])
        answer_callback(cb["id"], "生成面試問題⋯⋯")
        threading.Thread(target=handle_job_questions, args=(job_id,), daemon=True).start()

    elif data.startswith("job_tips_"):
        job_id = int(data[len("job_tips_"):])
        answer_callback(cb["id"], "生成 Key Tips⋯⋯")
        threading.Thread(target=handle_job_tips, args=(job_id,), daemon=True).start()

    elif data.startswith("job_updatestatus_"):
        job_id = int(data[len("job_updatestatus_"):])
        answer_callback(cb["id"])
        handle_update_status_menu(job_id)

    elif data.startswith("job_status_"):
        # format: job_status_{id}_{status}
        rest   = data[len("job_status_"):]
        parts  = rest.split("_", 1)
        job_id = int(parts[0])
        status = parts[1].replace("_", " ")
        jobs   = load_jobs()
        for j in jobs:
            if j["id"] == job_id:
                j["status"] = status
                break
        save_jobs(jobs)
        answer_callback(cb["id"], f"✅ 已更新為 {status}")
        send_telegram(
            f"{STATUS_EMOJI.get(status, '📝')} *{status}* 已更新！\n\n查看所有申請：/listjobs",
        )

    elif data.startswith("job_negotiate_"):
        job_id = int(data[len("job_negotiate_"):])
        jobs   = load_jobs()
        job    = next((j for j in jobs if j["id"] == job_id), None)
        if not job:
            answer_callback(cb["id"], "搵唔到呢個申請記錄")
        else:
            answer_callback(cb["id"])
            offer_details = f"職位：{job['role']}\n公司：{job['company']}\n（其他 package 詳情可以喺對話中補充）"
            save_session({"state": "negotiate_session", "offer_details": offer_details, "round_num": 0, "history": []})
            send_telegram(
                f"🤝 即將同 *{job['company']}* 嘅 HR 傾 *{job['role']}* 嘅薪酬。\n\n打你想講嘅第一句：",
                reply_markup={"inline_keyboard": [[{"text": "🏁 結束談判", "callback_data": "negotiate_end"}]]}
            )

    elif data == "negotiate_cancel":
        answer_callback(cb["id"], "已取消")
        clear_session()
        send_telegram("已取消談判練習。")

    elif data == "negotiate_end":
        answer_callback(cb["id"])
        session = load_session()
        clear_session()
        threading.Thread(target=_handle_negotiate_summary, args=((session or {}).get("history", []),), daemon=True).start()

    elif data == "debrief_job_skip":
        answer_callback(cb["id"])
        save_session({"state": "debrief_input", "job_info": None})
        send_telegram(DEBRIEF_PROMPT)

    elif data.startswith("debrief_job_"):
        answer_callback(cb["id"])
        job_id = int(data[len("debrief_job_"):])
        jobs   = load_jobs()
        job    = next((j for j in jobs if j["id"] == job_id), None)
        save_session({"state": "debrief_input", "job_info": job})
        send_telegram(DEBRIEF_PROMPT)

    elif data.startswith("job_practice_"):
        job_id = int(data[len("job_practice_"):])
        jobs   = load_jobs()
        job    = next((j for j in jobs if j["id"] == job_id), None)
        if job:
            answer_callback(cb["id"], f"生成 {job['company']} 面試場景⋯⋯")
            # 用 job 嘅 role 同 company 作 context 練習
            profile = load_profile()
            profile["_job_context"] = f"{job['company']} — {job['role']}"
            threading.Thread(
                target=start_practice,
                kwargs={"force_industry": profile.get("industry")},
                daemon=True,
            ).start()
        else:
            answer_callback(cb["id"], "找不到申請記錄")


def _send_setup_done(profile: dict, extra: str = ""):
    job   = profile.get("job_title", "未設定")
    ind   = profile.get("industry",  "未設定")
    mbti  = profile.get("mbti",      "未設定")
    send_telegram(
        f"✅ 設定完成！\n\n"
        f"🎯 目標職位：{job}\n"
        f"🏭 行業：{ind}\n"
        f"🧠 MBTI：{mbti}"
        f"{extra}\n\n"
        f"用 /practice 開始你嘅面試練習！",
        reply_markup={"inline_keyboard": [[
            {"text": "🎯 立即練習", "callback_data": "practice_new"},
        ]]}
    )


# ── 指令處理 ──────────────────────────────────────────────────────

def cmd(text: str, command: str) -> bool:
    return text == command or text.startswith(command + "@") or text.startswith(command + " ")


def _is_job_url(text: str) -> bool:
    t = text.strip()
    return bool(re.match(r"https?://\S+", t)) and len(t.split()) == 1


def handle_message(text: str):
    # ── URL 全自動 add：fetch → DeepSeek 抽取 → 直接 save ──────────
    addjob = load_addjob_session()
    if not addjob and _is_job_url(text):
        send_telegram("🔗 收到職位 link，正在讀取職位資料...")
        threading.Thread(target=_auto_add_job_from_url, args=(text.strip(),), daemon=True).start()
        return

    # ── Add Job session ──
    if addjob and not text.startswith("/"):
        state = addjob.get("state")

        if state == "addjob_company":
            addjob["company"] = text.strip()
            addjob["state"]   = "addjob_role"
            save_addjob_session(addjob)
            send_telegram(f"✅ 公司：{text.strip()}\n\n💼 職位名稱係？（例如：Admissions Officer）")
            return

        if state == "addjob_role":
            addjob["role"] = text.strip()
            # 如果 link 已預填（由 URL quick-add 觸發），直接 save，跳過 JD + link 步驟
            if addjob.get("link"):
                from datetime import date as _date
                jobs   = load_jobs()
                new_id = (max(j["id"] for j in jobs) + 1) if jobs else 1
                jobs.append({
                    "id":           new_id,
                    "company":      addjob.get("company", ""),
                    "role":         addjob["role"],
                    "jd":           "",
                    "link":         addjob["link"],
                    "applied_date": str(_date.today()),
                    "status":       "Applied",
                })
                save_jobs(jobs)
                clear_addjob_session()
                send_telegram(
                    f"✅ 已新增申請！\n\n"
                    f"🏢 {addjob.get('company')} — {addjob['role']}\n"
                    f"📅 {_date.today()}  |  狀態：Applied\n\n"
                    f"用 /listjobs 更新狀態或生成面試問題。"
                )
                return
            addjob["state"] = "addjob_jd"
            save_addjob_session(addjob)
            send_telegram(
                f"✅ 職位：{text.strip()}\n\n📄 貼上 JD 內容（可以直接 copy paste）\n\n輸入 /skip 跳過",
            )
            return

        if state == "addjob_jd":
            addjob["jd"]    = "" if text == "/skip" else text.strip()
            addjob["state"] = "addjob_link"
            save_addjob_session(addjob)
            send_telegram("✅ JD 已儲存！\n\n🔗 Job link（例如 JobsDB URL），輸入 /skip 跳過：")
            return

        if state == "addjob_link":
            addjob["link"] = "" if text == "/skip" else text.strip()
            # Save the job
            from datetime import date as _date
            jobs   = load_jobs()
            new_id = (max(j["id"] for j in jobs) + 1) if jobs else 1
            new_job = {
                "id":           new_id,
                "company":      addjob.get("company", ""),
                "role":         addjob.get("role", ""),
                "jd":           addjob.get("jd", ""),
                "link":         addjob.get("link", ""),
                "applied_date": str(_date.today()),
                "status":       "Applied",
            }
            jobs.append(new_job)
            save_jobs(jobs)
            clear_addjob_session()
            send_telegram(
                f"✅ *申請 #{new_id} 已記錄！*\n\n"
                f"🏢 {new_job['company']}\n"
                f"💼 {new_job['role']}\n"
                f"📅 申請日：{new_job['applied_date']}\n"
                f"📊 狀態：Applied\n\n"
                f"用 /listjobs 查看所有申請",
                reply_markup={"inline_keyboard": [[
                    {"text": "❓ 生成面試問題", "callback_data": f"job_q_{new_id}"},
                    {"text": "💼 Key Tips",     "callback_data": f"job_tips_{new_id}"},
                ]]}
            )
            return

    # Setup session 中：非指令訊息 = 填寫職位 / 公司
    setup = load_setup_session()
    if setup and not text.startswith("/"):
        state = setup.get("state")

        if state == "setup_industry_custom":
            profile = load_profile()
            profile["industry"] = text.strip()
            save_profile(profile)
            save_setup_session({"state": "setup_jobtitle"})
            send_telegram(f"✅ 行業：{text.strip()}\n\n🎯 你嘅目標職位係？（例如：UX Designer）")
            return

        if state == "setup_jobtitle":
            profile = load_profile()
            profile["job_title"] = text.strip()
            save_profile(profile)
            save_setup_session({"state": "setup_salary"})
            send_telegram("💰 仲有最後一步！你目標月薪期望大概係幾多？（例如：38000 或 38K）")
            return

        if state == "setup_salary":
            profile = load_profile()
            expected_salary = parse_salary_input(text)
            profile["expected_salary"] = expected_salary
            profile["salary_currency"] = "HKD"
            save_profile(profile)
            send_telegram("💰 分析緊薪酬市場數據⋯⋯")
            threading.Thread(
                target=_finish_salary_step,
                args=(profile, expected_salary),
                daemon=True,
            ).start()
            return

        if state == "setup_mbti":
            # 手動打 MBTI（e.g. "INTJ"）
            mbti_input = text.strip().upper()
            if mbti_input in MBTI_COACHING:
                profile = load_profile()
                profile["mbti"] = mbti_input
                save_profile(profile)
                clear_setup_session()
                coaching = MBTI_COACHING[mbti_input]
                _send_setup_done(profile, extra=f"\n\n💡 {mbti_input} 面試特點：{coaching['watch_out']}")
            else:
                send_telegram("請從鍵盤揀選 MBTI，或者打「跳過」")
            return

    # Practice / Review session 中
    session = load_session()

    if session and session.get("state") == "waiting_review" and not text.startswith("/"):
        clear_session()
        profile = load_profile()
        send_telegram("🔍 AI 分析緊你嘅面試對話，稍等⋯⋯")
        threading.Thread(
            target=lambda: send_telegram(analyze_conversation(text, profile)),
            daemon=True,
        ).start()
        return

    if session and session.get("state") == "waiting_response" and not text.startswith("/"):
        log.info(f"練習回應: {text[:60]!r}")
        threading.Thread(target=handle_user_response, args=(text,), daemon=True).start()
        return

    if session and session.get("state") == "negotiate_start" and not text.startswith("/"):
        session["offer_details"] = text.strip()
        session["round_num"]     = 0
        session["history"]       = []
        session["state"]         = "negotiate_session"
        save_session(session)
        send_telegram(
            "✅ Offer 已記錄！而家你可以開始同 HR 傾，打你想講嘅第一句：",
            reply_markup={"inline_keyboard": [[{"text": "🏁 結束談判", "callback_data": "negotiate_end"}]]}
        )
        return

    if session and session.get("state") == "negotiate_session":
        if text.strip() == "結束" or cmd(text, "/negotiate"):
            clear_session()
            threading.Thread(target=_handle_negotiate_summary, args=(session.get("history", []),), daemon=True).start()
            return
        threading.Thread(target=_handle_negotiate_turn, args=(session, text.strip()), daemon=True).start()
        return

    if session and session.get("state") == "debrief_input" and not text.startswith("/"):
        job_info = session.get("job_info")
        clear_session()
        threading.Thread(target=_handle_debrief_result, args=(job_info, text.strip()), daemon=True).start()
        return

    # ── 指令 ──

    if cmd(text, "/addjob"):
        handle_addjob_start()
        return

    if cmd(text, "/listjobs") or cmd(text, "/jobs"):
        handle_listjobs()
        return

    if cmd(text, "/setup"):
        clear_setup_session()
        handle_setup_start(intro=False)
        return

    if cmd(text, "/help") or cmd(text, "/start"):
        profile = load_profile()
        if not profile.get("job_title") and not profile.get("industry"):
            send_telegram(
                "👋 歡迎！AI 面試教練幫你針對練習，提升面試表現。\n\n"
                "首先設定你嘅背景，練習會更貼近你嘅實際情況：",
                reply_markup={"inline_keyboard": [
                    [{"text": "📄 上傳 CV（PDF / DOCX）", "callback_data": "onboard_cv"}],
                    [{"text": "✏️ 自己輸入背景",          "callback_data": "onboard_manual"}],
                ]}
            )
            return
        send_telegram(
            "🎓 AI 面試教練\n\n"
            "━━ 求職追蹤 ━━\n"
            "/addjob — 新增求職申請記錄\n"
            "/listjobs — 查看所有申請 + 狀態（可生成題目 / Tips / 練習）\n"
            "📊 CRM Kanban：https://sales-trainer-jatucpwszxyvoq5kpt7bav.streamlit.app\n\n"
            "━━ 面試練習 ━━\n"
            "/practice — 隨機面試練習\n"
            "/practice 初級／中級／高級 — 指定難度\n"
            "/drill — 針對特定題型練習\n"
            "/review — 貼真實面試答案，AI 分析失分點\n"
            "/negotiate — 薪酬談判 role-play\n"
            "/debrief — 面試後覆盤分析\n\n"
            "━━ 進度 ━━\n"
            "/stats — 查看進度報告\n"
            "/streak — 練習連續天數\n"
            "/tip — 今日面試技巧\n"
            "/setup — 更改目標職位 + MBTI\n"
            "/mystatus — 我的設定\n\n"
            "💡 每日練習 10 分鐘，面試表現 30 日內明顯提升"
        )
        return

    if cmd(text, "/practice"):
        profile = load_profile()
        if not profile.get("industry") and not profile.get("job_title"):
            handle_setup_start(intro=True)
            return

        parts = text.split(maxsplit=1)
        extra = parts[1].strip() if len(parts) > 1 else None
        diff  = extra if extra in DIFFICULTY_LEVELS else None

        send_telegram("🎯 生成面試場景⋯⋯")
        threading.Thread(
            target=start_practice,
            kwargs={"difficulty": diff},
            daemon=True,
        ).start()
        return

    if cmd(text, "/drill"):
        handle_drill_menu()
        return

    if cmd(text, "/review"):
        save_session({"state": "waiting_review"})
        send_telegram(
            "📋 貼上你嘅真實面試問答記錄，AI 幫你分析失分點同改善方向：\n\n"
            "（格式：面試官問：xxx\n我答：xxx）"
        )
        return

    if cmd(text, "/negotiate"):
        save_session({"state": "negotiate_start"})
        send_telegram(
            "🤝 薪酬談判練習\n\n貼你收到嘅 offer details（職位、公司、提供月薪、其他 package）：",
            reply_markup={"inline_keyboard": [[{"text": "❌ 取消", "callback_data": "negotiate_cancel"}]]}
        )
        return

    if cmd(text, "/debrief"):
        jobs = load_jobs()
        if jobs:
            keyboard = [
                [{"text": f"{STATUS_EMOJI.get(j.get('status','Applied'),'')} {j['company']} — {j['role']}",
                  "callback_data": f"debrief_job_{j['id']}"}]
                for j in jobs
            ]
            keyboard.append([{"text": "⏭️ 跳過，唔連結特定工", "callback_data": "debrief_job_skip"}])
            save_session({"state": "debrief_job_select"})
            send_telegram("🎙️ 係邊份工嘅面試？", reply_markup={"inline_keyboard": keyboard})
        else:
            save_session({"state": "debrief_input", "job_info": None})
            send_telegram(DEBRIEF_PROMPT)
        return

    if cmd(text, "/stats"):
        handle_stats()
        return

    if cmd(text, "/streak"):
        data   = load_stats()
        streak = data.get("streak", {})
        total  = data.get("total_sessions", 0)
        count  = streak.get("count", 0)
        emoji  = "🔥" if count >= 7 else "💪" if count >= 3 else "🌱"
        send_telegram(
            f"{emoji} 連續練習：{count} 日\n"
            f"總練習次數：{total} 次\n\n"
            f"每日練習，面試信心係咁升！"
        )
        return

    if cmd(text, "/tip"):
        handle_tip()
        return

    if cmd(text, "/mystatus"):
        profile = load_profile()
        data    = load_stats()
        total   = data.get("total_sessions", 0)
        streak  = data.get("streak", {}).get("count", 0)
        mbti    = profile.get("mbti", "未設定")
        coaching_note = ""
        if profile.get("mbti") and profile["mbti"].upper() in MBTI_COACHING:
            c = MBTI_COACHING[profile["mbti"].upper()]
            coaching_note = f"\n💡 你嘅面試盲點：{c['watch_out']}"
        send_telegram(
            f"⚙️ 我的設定\n\n"
            f"🎯 目標職位：{profile.get('job_title', '未設定')}\n"
            f"🏭 行業：{profile.get('industry', '未設定')}\n"
            f"🧠 MBTI：{mbti}{coaching_note}\n\n"
            f"📊 總練習次數：{total}\n"
            f"🔥 連續天數：{streak}\n\n"
            f"更改設定：/setup"
        )
        return


# ── 主循環 ────────────────────────────────────────────────────────

def poll():
    register_commands()
    log.info(f"AI 面試教練 Bot 啟動 {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    url    = f"https://api.telegram.org/bot{os.getenv('TELEGRAM_BOT_TOKEN')}/getUpdates"
    offset = None

    while True:
        try:
            params = {"timeout": 30, "allowed_updates": ["message", "callback_query"], "offset": offset}
            resp   = requests.get(url, params=params, timeout=35)
            data   = resp.json()

            for update in data.get("result", []):
                offset = update["update_id"] + 1
                if "callback_query" in update:
                    handle_callback(update["callback_query"])
                elif "message" in update:
                    document = update["message"].get("document")
                    text     = update["message"].get("text", "").strip()
                    if document:
                        log.info(f"收到 document: {document.get('file_name','')!r}")
                        threading.Thread(target=handle_document, args=(document,), daemon=True).start()
                    elif text:
                        log.info(f"收到: {text[:60]!r}")
                        handle_message(text)

        except Exception as e:
            log.error(f"Poll error: {e}")
            time.sleep(5)


if __name__ == "__main__":
    poll()
