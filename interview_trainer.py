"""AI 面試教練引擎：題目池、場景 DNA、MBTI coaching、AI 評估。"""

import random
import re
import os
from openai import OpenAI
from pathlib import Path
from dotenv import load_dotenv

BASE_DIR = Path(__file__).parent
load_dotenv(BASE_DIR / ".env")

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")

from utils import load_recent_dna, save_recent_dna, send_telegram


# ── 面試題型池（weight 越高越常出現）──────────────────────────────
QUESTION_TYPES = [
    {
        "name": "行為題 STAR",
        "weight": 4,
        "category": "behavioral",
        "example_q": "請講一個你喺工作中遇到衝突，最後點解決嘅例子。",
        "tip": "用 STAR 結構：Situation（背景）→ Task（任務）→ Action（行動）→ Result（結果）。Result 要有具體數字或成果，唔好只講過程。",
    },
    {
        "name": "優缺點自我評估",
        "weight": 3,
        "category": "self-awareness",
        "example_q": "你覺得自己最大嘅弱點係咩？",
        "tip": "講真實嘅弱點，但要展示你已經意識到並主動改善。唔好講「我太完美主義」呢類套路答案，面試官聽厭左。",
    },
    {
        "name": "職涯規劃",
        "weight": 3,
        "category": "motivation",
        "example_q": "你五年後想係邊度？點解揀我哋公司？",
        "tip": "答案要同目標公司嘅發展方向有 alignment。展示你對呢個行業有長遠承諾，唔係純粹跳板。",
    },
    {
        "name": "壓力處理",
        "weight": 2,
        "category": "resilience",
        "example_q": "你試過同時面對多個 deadline，你係點應對嘅？",
        "tip": "展示你有系統地處理壓力：優先排序、溝通期望值、保持冷靜。要有具體例子，唔係講理論。",
    },
    {
        "name": "團隊合作衝突",
        "weight": 3,
        "category": "interpersonal",
        "example_q": "你試過同唔同意你睇法嘅同事合作，你係點處理嘅？",
        "tip": "唔好話同事錯或自己永遠啱。重點係展示你識尊重不同意見、有方法找到共識，同時堅持原則。",
    },
    {
        "name": "領導力",
        "weight": 2,
        "category": "leadership",
        "example_q": "講一次你帶領團隊克服困難嘅經歷。",
        "tip": "唔一定要係 manager 先有領導力故事。可以係 project lead、mentor junior、或者主動推動改變嘅例子。",
    },
    {
        "name": "為何揀我哋",
        "weight": 3,
        "category": "motivation",
        "example_q": "你點解想加入我哋公司？係咩吸引到你？",
        "tip": "做 research！講出公司具體嘅產品、文化、近期動態。唔好講「公司大」「薪酬好」呢類答案，要展示你真係有興趣。",
    },
    {
        "name": "薪酬期望談判",
        "weight": 2,
        "category": "negotiation",
        "example_q": "你嘅薪酬期望係幾多？",
        "tip": "事先做 market research，提出一個合理範圍而唔係死價。可以先問對方 budget range，或者用「根據市場同我嘅經驗，我期望係 X–Y」框架。",
    },
    {
        "name": "情境判斷",
        "weight": 2,
        "category": "situational",
        "example_q": "如果你發現同事違反公司政策，你會點做？",
        "tip": "呢類題冇標準答案，但要展示你有原則、會按情況判斷、唔會盲目遵從或告密。答案要平衡公司利益同人際關係。",
    },
    {
        "name": "技術／專業知識",
        "weight": 2,
        "category": "technical",
        "example_q": "你係點keep up最新行業趨勢嘅？舉個例你近期學到嘅野。",
        "tip": "展示你有持續學習嘅習慣。最好可以提到具體嘅書、課程、會議、或者你係點apply新知識落工作。",
    },
]


# ── 面試官性格池 ───────────────────────────────────────────────────
INTERVIEWER_PERSONAS = [
    {"name": "友善 HR",       "desc": "態度親切，想了解你呢個人，但會問深入嘅 culture fit 問題"},
    {"name": "冷峻技術面試官", "desc": "直接、少廢話，只關心你有冇料到，唔係好在意你答嘢係咪流暢"},
    {"name": "壓力測試型",    "desc": "故意質疑你嘅答案，問追問題，測試你喺壓力下係咪冷靜"},
    {"name": "C-level 高管",  "desc": "大局思維，問題係戰略性嘅，唔係細節，想知你點 think big"},
    {"name": "Panel 多人面試", "desc": "三個面試官輪流發問，風格不一，你要 manage 唔同期望"},
    {"name": "沉默考驗型",    "desc": "你講完之後唔係好即時回應，用沉默試試你係咪會 over-explain 或者動搖"},
]


# ── 面試輪次池 ────────────────────────────────────────────────────
INTERVIEW_ROUNDS = [
    "電話初篩，對方語氣輕鬆但問題直接，大約 15 分鐘",
    "第一輪 HR 面試，視像會議，公司想了解你嘅背景同動機",
    "技術面試，面試官係你未來直屬上司，問得好具體",
    "Case Study 環節，你需要分析一個商業問題並即場回應",
    "最終 Management Round，高層親自面試，氣氛正式",
    "Offer 前最後一關，HR 跟進薪酬期望同入職時間",
]


# ── 行業池（14種）────────────────────────────────────────────────
INDUSTRIES = [
    "金融／投資銀行",
    "科技／軟件開發",
    "市場營銷／廣告",
    "管理諮詢",
    "零售／酒店管理",
    "初創公司",
    "醫療／藥劑／健康科技",
    "法律／合規",
    "人力資源／招聘",
    "教育／培訓",
    "物流／供應鏈",
    "地產／建築工程",
    "傳媒／娛樂／創意",
    "政府／NGO／公共服務",
]


# ── 難度 ──────────────────────────────────────────────────────────
DIFFICULTY_LEVELS = {
    "初級": "面試官態度友善，問題直接，適合練習基本結構",
    "中級": "面試官有追問，需要展示深度，部分問題有陷阱",
    "高級": "面試官態度強硬，質疑你嘅答案，需要高度自信同靈活應對",
}

SCORE_LABELS = {
    1: "❌ 需要改善",
    2: "⚠️ 有進步空間",
    3: "✅ 不錯",
    4: "🌟 出色",
}


# ── MBTI Coaching（16種）─────────────────────────────────────────
MBTI_COACHING = {
    "INTJ": {
        "label": "INTJ — 策略家",
        "strengths": "思路清晰、答嘢有深度、唔容易被 pressure",
        "watch_out": "聽落太冷、缺乏 warmth，面試官覺得你唔好相處；答嘢太長太理論，缺少具體故事",
        "tip": "答每條題時先用一句 anchor 你嘅核心 point，再給例子。練習面試前段加一句 small talk 或 genuine 嘅讚美，soften 第一印象。",
    },
    "INTP": {
        "label": "INTP — 學者",
        "strengths": "分析能力強，答 technical 題有深度，思考全面",
        "watch_out": "答嘢太 open-ended、兜圈、面試官唔知你要表達咩；而且容易說「依我睇有好多可能性……」而唔係 commit 一個答案",
        "tip": "每條答案先說結論，再解釋。練習：答完之後問自己「我嘅主要 point 係咩？」先 check 係咪清楚。",
    },
    "ENTJ": {
        "label": "ENTJ — 指揮官",
        "strengths": "自信、有領袖氣質、答領導力題自然有說服力",
        "watch_out": "聽落太強勢、甚至傲慢；面試官覺得你唔係 team player；唔夠展示 empathy",
        "tip": "刻意在答案中加入「我哋 team」而唔係淨係「我」。問追問題時用「我想更了解……」而唔係直接 challenge 面試官。",
    },
    "ENTP": {
        "label": "ENTP — 辯論家",
        "strengths": "思維靈活、反應快、答 situational 題有創意",
        "watch_out": "答嘢散、跳太快、結構唔夠；面試官唔確定你係咪 reliable；可能顯得「諗太多、做太少」",
        "tip": "嚴格用 STAR 結構。每個答案 2 分鐘內完成。練習收尾：「所以最後結果係……」確保有 landing。",
    },
    "INFJ": {
        "label": "INFJ — 提倡者",
        "strengths": "答 values 同 purpose 嘅題特別有感染力，genuine、有深度",
        "watch_out": "太內向、答嘢太「深」令面試官聽唔明；唔擅長 sell 自己，說成就時太謙虛",
        "tip": "把你嘅 impact 量化：「因為我嘅建議，team 節省咗 X 小時」。練習大聲說出自己嘅成就，唔係謙虛係事實。",
    },
    "INFP": {
        "label": "INFP — 調停者",
        "strengths": "真誠、有原則、文化 fit 問題答嘢感人",
        "watch_out": "太理想化、答嘢飄；唔夠 confident；講緊自己嘅成就時會 downplay",
        "tip": "練習用 assertive 語氣：「我主導咗呢個項目」而唔係「我有份參與」。感受可以講，但要配合具體事實。",
    },
    "ENFJ": {
        "label": "ENFJ — 主人公",
        "strengths": "有 charisma、答 teamwork 同 leadership 題有感染力、面試官喜歡你",
        "watch_out": "太 people-pleasing，答嘢為咗面試官而答，而唔係真實；有時唔夠 assertive 表達自己立場",
        "tip": "練習回答有爭議性嘅情境題時 hold 住自己嘅立場，唔好見面試官皺眉就即刻改口。",
    },
    "ENFP": {
        "label": "ENFP — 競選者",
        "strengths": "熱情、有創意、storytelling 自然，面試官記得你",
        "watch_out": "答嘢太長太散、結構唔夠；容易講到去唔關題目事嘅野；面試官唔確定你係咪 detail-oriented",
        "tip": "每條答案開頭就 state structure：「我分三點講……」。限制自己每個 STAR 例子最多 90 秒，練習計時。",
    },
    "ISTJ": {
        "label": "ISTJ — 物流師",
        "strengths": "答嘢有條理、可靠、細節豐富，面試官信任你",
        "watch_out": "聽落太保守、太 rule-follower；唔夠展示創新或 adaptability；答嘢太 dry，缺少 energy",
        "tip": "每隔幾條題主動加一個「我哋係咁做，但我覺得可以改進係……」展示你唔只係執行者，仲有思考。",
    },
    "ISFJ": {
        "label": "ISFJ — 守衛者",
        "strengths": "reliable、有 empathy、teamwork 題有真誠感",
        "watch_out": "唔夠 sell 自己，把成功歸功於 team 而唔係自己；唔擅長談薪酬；怕衝突所以情境題答得太保守",
        "tip": "用「我嘅貢獻係……」而唔係「我哋做咗……」。談薪酬時練習用數據支持你嘅期望，而唔係純粹講感覺。",
    },
    "ESTJ": {
        "label": "ESTJ — 總經理",
        "strengths": "自信、有 authority、答 leadership 同 process 題清晰有力",
        "watch_out": "聽落太 rigid、唔夠 open-minded；答情境題傾向「我就係咁決定」而缺少考慮不同 stakeholders",
        "tip": "答每個決定時加一句：「但我同時考慮到……」展示你有全局觀，而唔係只係執行自己意志。",
    },
    "ESFJ": {
        "label": "ESFJ — 執政官",
        "strengths": "熱情、答 culture fit 同 teamwork 題特別好、面試官第一印象佳",
        "watch_out": "太在意面試官反應、容易改口；唔夠展示獨立思考；壓力下容易緊張",
        "tip": "練習喺面試官 silent 或者皺眉時唔好 panic，繼續原有答案。Silence 唔代表你答錯，係面試官思考緊。",
    },
    "ISTP": {
        "label": "ISTP — 鑑賞家",
        "strengths": "答 technical 同 problem-solving 題有深度，冷靜不亂",
        "watch_out": "答嘢太簡短，唔夠 elaborate；面試官問「點解」時唔知點展開；缺少情感溫度",
        "tip": "每條答案練習「強制拉長」：答完之後加一句「呢個經驗令我學到……」。目標每個 STAR 答案不少於 60 秒。",
    },
    "ISFP": {
        "label": "ISFP — 探險家",
        "strengths": "真誠、有獨特創意、答 passion 同 values 題有感染力",
        "watch_out": "太 low-key、唔夠 sell 自己；面對 pressure 型面試官會退縮；結構散",
        "tip": "面試前寫低 3 個你最自豪嘅成就，練習用一句說清楚每個。面試時如果緊張，深呼吸然後問自己「我想比佢知乜嘢？」",
    },
    "ESTP": {
        "label": "ESTP — 企業家",
        "strengths": "反應快、自信、答 negotiation 同 high-pressure 題自然",
        "watch_out": "答嘢太即興、唔夠深度；面試官覺得你冇認真諗；長期規劃問題答得比較弱",
        "tip": "呢種面試要事先準備多 D。每個題型最少想好一個 go-to 例子。「即興發揮」喺面試係風險。",
    },
    "ESFP": {
        "label": "ESFP — 表演者",
        "strengths": "充滿能量、storytelling 自然生動、面試官喜歡你呢個人",
        "watch_out": "答嘢太 surface level；容易離題；面試官唔確定你係咪夠 serious 同 detail-oriented",
        "tip": "每條答案結尾加一句總結：「所以呢件事令我明白到……」 Anchor 住你嘅 point，唔好飄走。",
    },
}

# MBTI 維度快速 tips（用於 prompt 補充）
_MBTI_DIM_TIPS = {
    "I": "性格偏內向，記得主動展開答案，面試官唔會主動問你挖更多，你要自己 volunteer 資訊。",
    "E": "性格偏外向，注意唔好搶話或者答太長，留空間比面試官消化。",
    "N": "傾向宏觀思維，記得加具體例子支持你每個 point，唔好只講概念。",
    "S": "答嘢細節豐富，記得每隔一陣 zoom out 講下 big picture 同你嘅思路。",
    "T": "偏邏輯分析，記得喺答案加入情感智慧元素，例如點理解 stakeholder 嘅感受。",
    "F": "偏重人際感受，記得自信說出自己嘅成就同立場，唔係謙虛係必要嘅 self-advocacy。",
    "J": "有計劃有結構，記得展示你嘅靈活性同 adaptability，唔係死板遵守計劃。",
    "P": "靈活適應，記得展示你嘅計劃性同可靠度，面試官想知你係咪 follow through。",
}


def get_mbti_context(mbti: str) -> str:
    """生成 MBTI coaching context 供 prompt 用。"""
    mbti = mbti.upper().strip()
    if mbti not in MBTI_COACHING:
        return ""

    coaching = MBTI_COACHING[mbti]
    dims = list(mbti)  # e.g. ['I','N','T','J']
    dim_tips = [_MBTI_DIM_TIPS.get(d, "") for d in dims if _MBTI_DIM_TIPS.get(d)]

    return f"""
【學員 MBTI：{coaching['label']}】
優勢：{coaching['strengths']}
面試常見盲點：{coaching['watch_out']}
針對性建議：{coaching['tip']}
維度提示：{' '.join(dim_tips)}"""


# ── DNA 抽選（防重複）────────────────────────────────────────────

def _pick_fresh_dict(pool: list, key: str, recent: dict, window: int = 4):
    used  = set(recent.get(key, []))
    fresh = [x for x in pool if x["name"] not in used] or pool
    weights = [x.get("weight", 1) for x in fresh]
    return random.choices(fresh, weights=weights, k=1)[0]

def _pick_fresh_str(pool: list, key: str, recent: dict, window: int = 4):
    used  = set(recent.get(key, []))
    fresh = [x for x in pool if x not in used] or pool
    return random.choice(fresh)

def _update_recent(recent: dict, key: str, val: str, window: int = 4):
    lst = recent.get(key, [])
    if val in lst:
        lst.remove(val)
    lst.append(val)
    recent[key] = lst[-window:]


def pick_scenario_dna(force_qtype: str = None, force_industry: str = None, difficulty: str = None) -> dict:
    """抽取場景 DNA，支援強制指定題型或行業，自動防重複。"""
    try:
        recent = load_recent_dna()
    except Exception:
        recent = {}

    # Question type
    if force_qtype:
        qtype = next((x for x in QUESTION_TYPES if force_qtype in x["name"]), None) \
                or _pick_fresh_dict(QUESTION_TYPES, "qtype", recent)
    else:
        qtype = _pick_fresh_dict(QUESTION_TYPES, "qtype", recent)

    # Industry
    if force_industry:
        ind = next((x for x in INDUSTRIES if force_industry in x), None) \
              or _pick_fresh_str(INDUSTRIES, "industry", recent)
    else:
        ind = _pick_fresh_str(INDUSTRIES, "industry", recent)

    persona  = _pick_fresh_dict(INTERVIEWER_PERSONAS, "persona", recent)
    iv_round = _pick_fresh_str(INTERVIEW_ROUNDS, "iv_round", recent)
    diff     = difficulty or random.choices(
        list(DIFFICULTY_LEVELS.keys()), weights=[3, 4, 2], k=1
    )[0]

    _update_recent(recent, "qtype",    qtype["name"])
    _update_recent(recent, "industry", ind)
    _update_recent(recent, "persona",  persona["name"])
    _update_recent(recent, "iv_round", iv_round)
    try:
        save_recent_dna(recent)
    except Exception:
        pass

    return {
        "qtype":    qtype,
        "persona":  persona,
        "iv_round": iv_round,
        "industry": ind,
        "difficulty": diff,
    }


# ── 場景生成 ──────────────────────────────────────────────────────

def generate_scenario(force_qtype: str = None, force_industry: str = None, difficulty: str = None):
    """生成面試練習場景，返回 (展示文字, 場景 dict)。"""
    s = pick_scenario_dna(force_qtype, force_industry, difficulty)
    qtype   = s["qtype"]
    persona = s["persona"]

    display = (
        f"🎯 面試練習開始！\n"
        f"━━━━━━━━━━━━━━━━━━━━\n"
        f"🏭 行業：{s['industry']}\n"
        f"📍 輪次：{s['iv_round']}\n"
        f"👤 面試官：{persona['name']} — {persona['desc']}\n"
        f"📝 題型：{qtype['name']}\n"
        f"⚡ 難度：{s['difficulty']}\n"
        f"━━━━━━━━━━━━━━━━━━━━\n\n"
        f"面試官問：\n「{qtype['example_q']}」\n\n"
        f"👇 請輸入你嘅回答（廣東話／英文均可）"
    )
    return display, s


# ── AI 評估回應 ────────────────────────────────────────────────────

def evaluate_response(user_response: str, scenario: dict, max_retries: int = 3, profile: dict = None) -> str:
    """用 DeepSeek 評估面試回答，加入 MBTI coaching。"""
    ai_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
    qtype   = scenario["qtype"]
    persona = scenario["persona"]
    profile = profile or {}

    # Profile context
    profile_ctx = ""
    if profile.get("job_title") or profile.get("industry") or profile.get("company"):
        profile_ctx = "\n【學員背景】\n"
        if profile.get("job_title"): profile_ctx += f"目標職位：{profile['job_title']}\n"
        if profile.get("industry"):  profile_ctx += f"目標行業：{profile['industry']}\n"
        if profile.get("company"):   profile_ctx += f"目標公司：{profile['company']}\n"
        if profile.get("exp_years"): profile_ctx += f"工作年資：{profile['exp_years']} 年\n"

    # MBTI context
    mbti_ctx = ""
    if profile.get("mbti"):
        mbti_ctx = get_mbti_context(profile["mbti"])

    prompt = f"""你係一個頂尖面試教練，同時扮演面試官角色。請評估以下面試練習。
{profile_ctx}{mbti_ctx}

【面試場景】
行業：{scenario['industry']}
輪次：{scenario['iv_round']}
難度：{scenario['difficulty']} — {DIFFICULTY_LEVELS[scenario['difficulty']]}
面試官性格：{persona['name']} — {persona['desc']}
題型：{qtype['name']}
面試官問：「{qtype['example_q']}」

【學員的回答】
{user_response}

【輸出格式——嚴格跟住以下結構，廣東話口語】

**面試官即時反應：**
（以呢個面試官嘅性格自然回應，1–2句。唔好咁快表示滿意，要保持真實壓力。）

**━━ 教練評分 ━━**
評分：X／4
（1=方向錯或令面試官有負面印象｜2=方向對但答法唔夠有力｜3=回答不錯，仲有提升空間｜4=出色，清晰有說服力有深度）

做得好的地方：（引用學員答案中具體嘅句子或點）

需要改善：（具體指出哪句話有問題，點改）
{f'【針對你係 {profile["mbti"]} 嘅特別提示】：{MBTI_COACHING[profile["mbti"].upper()]["tip"]}' if profile.get("mbti") and profile["mbti"].upper() in MBTI_COACHING else ""}

**━━ 最佳示範回應 ━━**
（用廣東話口語，自然真實，符合{scenario['difficulty']}難度，針對呢個面試官性格{('，以應徵' + profile['job_title'] + '嘅角度') if profile.get('job_title') else ''}）

**━━ 技巧重點 ━━**
{qtype['tip']}"""

    last_err = None
    for attempt in range(1, max_retries + 1):
        try:
            resp = ai_client.chat.completions.create(
                model="deepseek-chat",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.8,
                max_tokens=1400,
            )
            return resp.choices[0].message.content
        except Exception as e:
            last_err = e
            print(f"[evaluate_response] 第{attempt}次失敗：{e}")
    return f"⚠️ 評估失敗（已重試 {max_retries} 次）：{last_err}"


# ── 對話分析 ──────────────────────────────────────────────────────

def analyze_conversation(conversation: str, profile: dict) -> str:
    """分析用戶貼入嘅真實面試對話，返回詳細建議。"""
    ai_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
    job_title = profile.get("job_title", "")
    industry  = profile.get("industry", "")
    mbti      = profile.get("mbti", "")

    ctx_parts = []
    if industry:  ctx_parts.append(f"行業：{industry}")
    if job_title: ctx_parts.append(f"目標職位：{job_title}")
    ctx = "\n".join(ctx_parts) or "未指定"

    mbti_ctx = get_mbti_context(mbti) if mbti else ""

    prompt = f"""你係資深面試教練，分析以下真實面試對話並給出具體建議。

【求職背景】
{ctx}
{mbti_ctx}

【真實對話記錄】
{conversation}

【分析格式——廣東話口語】

**整體印象評分：X／4**

**最強嘅地方：**
（2–3點，引用具體句子）

**最需要改善：**
（2–3點，每點說明點改 + 示範改法）

**STAR 結構檢查：**
（呢個答案係咪有清晰嘅 Situation / Task / Action / Result？缺少咩？）

**下次面試要記住：**
（1–2句最重要嘅 takeaway）"""

    try:
        resp = ai_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=1200,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"⚠️ 分析失敗：{e}"


# ── 今日技巧 ──────────────────────────────────────────────────────

DAILY_TIPS = [
    "🎯 STAR 結構係面試嘅基本功。每個行為題都要有：S（背景）→ T（任務）→ A（你嘅行動）→ R（可量化成果）。",
    "💪 唔好怕沉默。面試官問完問題後，你可以說「俾我諗一秒先」，比即興亂答好得多。",
    "🔍 研究公司係必要功課。面試前搵 3 件近期新聞或公司動態，自然融入答案，展示你真係有興趣。",
    "📊 量化你嘅成就。「我改善咗效率」VS「我令 processing time 減少咗 30%」，後者有力 10 倍。",
    "🤝 面試係雙向嘅。你都係喺評估公司係唔係 fit 你。問問題唔係軟弱，係主動、係自信。",
    "💬 答「弱點」題嘅公式：真實弱點 + 你係點察覺到 + 你已採取咩行動改善。三部份缺一不可。",
    "🧠 準備 5–7 個核心故事，覆蓋：成功、失敗、衝突、領導、創新。大部分題目都可以用呢幾個故事 adapt。",
    "👁️ 視像面試提示：鏡頭要喺眼睛水平，背景整潔，著得正式。細節反映你對呢個機會嘅重視程度。",
    "🎤 練習大聲講出你嘅答案。喺腦海「諗到」同「說得流暢」係兩回事，一定要出聲練。",
    "💰 談薪酬唔好第一個報數。先問「你哋呢個職位嘅 budget range 係？」係合理嘅，唔係無禮。",
]

def get_daily_tip() -> str:
    from datetime import date
    idx = date.today().toordinal() % len(DAILY_TIPS)
    return DAILY_TIPS[idx]


# ── Job-specific AI functions ─────────────────────────────────────

def generate_job_questions(job: dict) -> str:
    """針對特定 JD 生成 6 個最可能出現的面試問題。"""
    ai_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
    company = job.get("company", "（未填）")
    role    = job.get("role",    "（未填）")
    jd      = job.get("jd",     "")

    prompt = f"""你係一個資深 HR，幫我針對以下職位生成 6 個最可能問嘅面試問題。

公司：{company}
職位：{role}
JD 內容：
{jd[:2000] if jd else "（未提供 JD）"}

【輸出格式——廣東話口語，直接輸出問題列表】

1. （行為題 / STAR 題）
2. （職位相關 competency 題）
3. （動機 / Why this company）
4. （壓力 / 衝突處理）
5. （領導 / 成就例子）
6. （未來規劃 / Career goal）

每條問題後加：
💡 回答重點：（一句建議）"""

    try:
        resp = ai_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=1000,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"⚠️ 生成失敗：{e}"


def generate_job_tips(job: dict) -> str:
    """針對特定 JD 生成 key talking points。"""
    ai_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
    company = job.get("company", "（未填）")
    role    = job.get("role",    "（未填）")
    jd      = job.get("jd",     "")

    prompt = f"""你係一個面試 coach，分析以下職位，提供面試前 key talking points。

公司：{company}
職位：{role}
JD：
{jd[:2000] if jd else "（未提供 JD）"}

【輸出格式——廣東話口語】

🎯 呢個職位最重視嘅 3 件事：
1.
2.
3.

💼 你應該重點 highlight 嘅經歷類型：
（2–3 條具體建議）

⚠️ 常見失分陷阱：
（1–2 條要避免嘅）

✨ 30 秒 Elevator Pitch 框架：
「我係...，有...年...經驗，喺...方面...，呢個職位吸引我因為...」"""

    try:
        resp = ai_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=800,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"⚠️ 生成失敗：{e}"


# ── Job URL 自動抽取 ──────────────────────────────────────────────

def extract_job_from_url(url: str) -> dict:
    """Fetch job URL → DeepSeek 抽出 company / role / jd。失敗返回空 dict。"""
    import requests as _req
    import json as _json

    ai_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

    try:
        resp = _req.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
        raw = resp.text[:6000]
    except Exception as e:
        print(f"[extract_job_from_url] fetch 失敗：{e}")
        return {}

    prompt = f"""你係一個 HR 助手。以下係一個求職頁面的 HTML 內容。
請抽取職位資訊，只輸出 JSON，唔要任何其他文字或 markdown：

{{
  "company": "公司名稱",
  "role": "職位名稱",
  "jd": "JD 內容撮要（最多 500 字，廣東話或英文均可）"
}}

如果某項資訊找不到，填空字串。

【頁面內容】
{raw}"""

    try:
        r = ai_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=600,
        )
        content = r.choices[0].message.content.strip()
        content = content.replace("```json", "").replace("```", "").strip()
        return _json.loads(content)
    except Exception as e:
        print(f"[extract_job_from_url] DeepSeek 失敗：{e}")
        return {}


# ── Resume 解析 ───────────────────────────────────────────────────

def parse_resume(resume_text: str) -> dict:
    """用 DeepSeek 分析 resume 文字，提取結構化 profile 資訊。"""
    import json as _json
    ai_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

    prompt = f"""你係一個 HR 專家，分析以下 resume 內容，提取結構化資訊。

【Resume 內容】
{resume_text[:6000]}

【輸出格式——只輸出 JSON，唔要任何其他文字，唔要 markdown code block】
{{
  "job_title": "最近或目標職位（例如：Marketing Manager）",
  "industry": "所屬行業（從以下選最近似：金融／投資銀行、科技／軟件開發、市場營銷／廣告、管理諮詢、零售／酒店管理、初創公司、醫療／藥劑／健康科技、法律／合規、人力資源／招聘、教育／培訓、物流／供應鏈、地產／建築工程、傳媒／娛樂／創意、政府／NGO／公共服務）",
  "exp_years": 估計工作年資數字（純數字，例如 5）,
  "current_company": "最近工作公司名稱",
  "key_skills": "3–5 個核心技能，逗號分隔",
  "education": "最高學歷，搵 Education / 學歷 / Qualification 等 section，唔到就搵大學名稱（例如：HKU 工商管理學士 2018）",
  "summary": "一句話描述用家背景（廣東話口語，用第二人稱「你」，例如：「你有5年教育行業經驗，專長係...」）"
}}

注意：education 係必填，就算只有一間大學名稱都要填，唔好填「未提供」。"""

    try:
        resp = ai_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=500,
        )
        content = resp.choices[0].message.content.strip()
        # 移除可能嘅 markdown fences
        content = content.replace("```json", "").replace("```", "").strip()
        return _json.loads(content)
    except Exception as e:
        print(f"[parse_resume] 失敗：{e}")
        return {}


# ── Cover Letter + Tailored CV 生成 ──────────────────────────────

def generate_cover_letter_from_jd(cv_text: str, jd_text: str, company: str, role: str) -> str:
    """根據 CV 全文 + JD，生成針對性 Cover Letter（英文，約 280 字）。"""
    ai_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

    prompt = f"""You are an expert cover letter writer for the Hong Kong job market with 10+ years experience.
Write a tailored, ATS-friendly cover letter for the following job application.

CRITICAL RULES:
1. Use ONLY facts, achievements, employers and skills that appear in the candidate's actual CV — NEVER fabricate experience, numbers, employers or titles.
2. Mirror the exact keywords and phrasing from the Job Description wherever the CV genuinely supports them (ATS optimisation).
3. Be specific — cite real achievements/metrics from the CV, not generic filler like "team player" or "hard-working".
4. Keep company names, job titles and dates exactly as written in the CV.
5. Match tone to the role and industry; confident but not boastful.

【Applicant's CV】
{cv_text[:3500]}

【Target Job】
Company: {company}
Role: {role}
Job Description:
{jd_text[:2500]}

【Output requirements】
- Length: 3 tight paragraphs, ~280 words
- Para 1: Who you are + the specific role/company + a hook tying your strongest relevant qualification to their need
- Para 2: 2–3 most relevant achievements from the CV, each explicitly matched to a JD requirement (use JD keywords)
- Para 3: Why this specific company (reference something concrete from the JD) + a confident call to action
- Plain text only, no markdown, ready to paste into an email
- Sign off with exactly: Yours sincerely,\n[Your Name]"""

    try:
        resp = ai_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.35,
            max_tokens=900,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"⚠️ 生成失敗：{e}"


def generate_tailored_cv_content(cv_text: str, jd_text: str, company: str, role: str) -> dict:
    """根據 JD 分析 CV，返回 tailored CV 各部份內容（dict），供 .docx 生成用。"""
    import json as _json
    ai_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

    prompt = f"""You are an expert ATS-optimised CV writer for the Hong Kong job market.
Tailor the applicant's CV for this specific job.

CRITICAL RULES:
1. Use ONLY information from the candidate's actual CV — NEVER fabricate companies, titles, dates, metrics or skills.
2. Keep all real job titles, company names and dates exactly as written in the CV.
3. Tailor the summary and every bullet to naturally match the JD's keywords (ATS optimisation) — but only where the CV genuinely supports it.
4. Core competencies: 8–12 items, mixing role-relevant hard skills AND tools/tech, drawn from the CV.
5. Each experience bullet should be achievement-oriented (action verb + what + result), and quantified whenever the CV provides numbers.
6. Do NOT mention degree/education or language ability inside the professional summary.

【Original CV】
{cv_text[:3500]}

【Target Job】
Company: {company}
Role: {role}
JD:
{jd_text[:2500]}

Output ONLY valid JSON (no markdown, no code block) with this exact structure:
{{
  "name": "full name from CV",
  "contact": "email | phone | location (one line)",
  "summary": "3-4 sentence professional summary tailored to this role, no education/language mention (English)",
  "core_competencies": ["8 to 12 skills mixing role-relevant hard skills and tools/tech"],
  "experience": [
    {{
      "company": "full company name",
      "role": "job title",
      "period": "date range",
      "bullets": ["4 achievement-oriented bullets, each matched to a JD requirement and quantified where possible"]
    }}
  ],
  "education": [
    {{"institution": "...", "degree": "...", "year": "..."}}
  ],
  "additional": "languages, certifications, tools (one line)"
}}"""

    try:
        resp = ai_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=2200,
        )
        content = resp.choices[0].message.content.strip()
        content = content.replace("```json", "").replace("```", "").strip()
        return _json.loads(content)
    except Exception as e:
        print(f"[generate_tailored_cv_content] 失敗：{e}")
        return {}


def build_cv_docx(cv_data: dict, company: str, role: str) -> bytes:
    """將 tailored CV dict 轉成 .docx bytes，可直接發 Telegram。"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()

    # 頁面設定
    section = doc.sections[0]
    section.top_margin    = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin   = Inches(0.8)
    section.right_margin  = Inches(0.8)

    BLUE = RGBColor(0x2E, 0x74, 0xB5)

    def heading(text, size=14, color=BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color
        return p

    def divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run("─" * 55)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    def section_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = BLUE

    def body(text, size=10, bold=False, space_after=2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        return p

    # ── Header ──
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_name.add_run(cv_data.get("name", ""))
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = BLUE

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = p_contact.add_run(cv_data.get("contact", ""))
    rc.font.size = Pt(9); rc.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p_contact.paragraph_format.space_after = Pt(2)

    # tailored tag
    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = p_tag.add_run(f"Tailored for: {role} @ {company}")
    rt.font.size = Pt(8); rt.italic = True
    rt.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    divider()

    # ── Summary ──
    section_title("Professional Summary")
    body(cv_data.get("summary", ""), space_after=4)

    # ── Core Competencies ──
    skills = cv_data.get("core_competencies", [])
    if skills:
        section_title("Core Competencies")
        cols = 3
        rows = (len(skills) + cols - 1) // cols
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        flat = skills + [""] * (rows * cols - len(skills))
        for i, sk in enumerate(flat):
            cell = table.cell(i // cols, i % cols)
            cell.text = f"• {sk}" if sk else ""
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
        doc.add_paragraph()

    # ── Experience ──
    section_title("Professional Experience")
    for exp in cv_data.get("experience", []):
        p_role = doc.add_paragraph()
        p_role.paragraph_format.space_after = Pt(0)
        r1 = p_role.add_run(exp.get("role", ""))
        r1.bold = True; r1.font.size = Pt(10)
        r2 = p_role.add_run(f"  |  {exp.get('company', '')}  |  {exp.get('period', '')}")
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        for b in exp.get("bullets", []):
            p_b = doc.add_paragraph(style="List Bullet")
            p_b.paragraph_format.space_after = Pt(1)
            run = p_b.add_run(b)
            run.font.size = Pt(9)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Education ──
    section_title("Education")
    for edu in cv_data.get("education", []):
        body(f"{edu.get('degree','')} — {edu.get('institution','')} {edu.get('year','')}", size=9)

    # ── Additional ──
    if cv_data.get("additional"):
        section_title("Additional")
        body(cv_data["additional"], size=9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── CV Health Score（本地計算，唔需要 AI）──────────────────────────

CV_ACTION_VERBS = [
    "led", "managed", "developed", "implemented", "increased", "reduced", "improved",
    "designed", "built", "launched", "optimized", "achieved", "delivered", "coordinated",
    "trained", "analyzed", "negotiated", "generated", "streamlined", "oversaw",
]

CV_KEYWORDS_GENERIC = [
    "stakeholder", "P&L", "KPI", "ROI", "SOP", "budget", "cross-functional",
    "project management", "data-driven", "client relations", "business development",
    "operations", "compliance", "reporting", "strategy",
]

CV_SECTION_PATTERNS = {
    "Summary":   r"summary|個人簡介|自我簡介|profile",
    "Work Exp":  r"work experience|employment|專業經驗|工作經驗|experience",
    "Skills":    r"skills|技能|專長|core competenc",
    "Education": r"education|學歷|教育背景",
}


def calculate_cv_health(cv_text: str) -> dict:
    """本地計算 CV Health Score（0-100，4 個維度：結構／量化成就／Action Verbs／關鍵詞），即時、唔需要 call AI。"""
    text = cv_text or ""
    lower = text.lower()

    # 1. 結構完整性（25）
    present, missing = [], []
    for name, pattern in CV_SECTION_PATTERNS.items():
        (present if re.search(pattern, lower) else missing).append(name)
    structure_score = max(0, 25 - len(missing) * 6)

    # 2. 量化成就（25）
    quant_count = len(re.findall(r"\d+(?:\.\d+)?%|\d{2,}", text))
    if quant_count >= 5:   quant_score = 25
    elif quant_count >= 3: quant_score = 18
    elif quant_count >= 1: quant_score = 10
    else:                  quant_score = 0

    # 3. Action Verbs（25）
    verb_count = sum(1 for v in CV_ACTION_VERBS if re.search(rf"\b{v}\b", lower))
    if verb_count >= 8:   verb_score = 25
    elif verb_count >= 5: verb_score = 18
    elif verb_count >= 3: verb_score = 10
    else:                 verb_score = 5

    # 4. 關鍵詞豐富度（25，對比 preset keyword list）
    matched_keywords = [k for k in CV_KEYWORDS_GENERIC if k.lower() in lower]
    keyword_score = round(len(matched_keywords) / len(CV_KEYWORDS_GENERIC) * 25)

    total = structure_score + quant_score + verb_score + keyword_score

    weak_points = []
    if quant_score < 18:
        weak_points.append(f"量化成就少（只有 {quant_count} 個數字／%）")
    if verb_score < 18:
        weak_points.append(f"強動詞少（只有 {verb_count} 個）")
    if keyword_score < 18:
        weak_points.append(f"行業關鍵詞唔夠（只有 {len(matched_keywords)} 個）")

    suggestions = []
    if missing:
        suggestions.append(f"補齊缺少嘅 {'／'.join(missing)} 部分")
    if quant_score < 18:
        suggestions.append("為成就加返具體數字或百分比")
    if verb_score < 18:
        suggestions.append("多用強動詞開頭（如 led／achieved／optimized）")
    suggestion = "；".join(suggestions[:2]) or "結構同內容都唔錯，可以再針對目標職位微調用詞。"

    return {
        "total": total,
        "structure_score": structure_score,
        "quant_score": quant_score,
        "verb_score": verb_score,
        "keyword_score": keyword_score,
        "sections_present": present,
        "sections_missing": missing,
        "weak_points": weak_points,
        "matched_keywords": matched_keywords,
        "suggestion": suggestion,
    }


def format_cv_health_message(health: dict) -> str:
    """將 calculate_cv_health() 嘅 dict 結果砌成 Telegram 訊息（webhook.py / bot_listener.py 共用）。"""
    lines = [f"📋 CV 基礎評分：{health['total']}/100", ""]
    if health["sections_present"]:
        lines.append(f"✅ 有：{'、'.join(health['sections_present'])}")
    if health["weak_points"]:
        lines.append(f"⚠️ 薄弱：{'；'.join(health['weak_points'])}")
    if health["sections_missing"]:
        lines.append(f"❌ 缺：{'、'.join(health['sections_missing'])}")
    lines.append("")
    lines.append(f"💡 建議：{health['suggestion']}")
    return "\n".join(lines)


def parse_salary_input(text: str) -> str:
    """從自由文字提取月薪數字，支援 38k／$38,000／HK$38000 等格式；解析唔到就原文返回。"""
    cleaned = (text or "").replace(",", "").upper()
    m = re.search(r"(\d+(?:\.\d+)?)\s*K\b", cleaned)
    if m:
        return str(int(float(m.group(1)) * 1000))
    m = re.search(r"\d{3,}", cleaned)
    if m:
        return m.group(0)
    return (text or "").strip()


# ── HK Salary Benchmark ──────────────────────────────────────────

def generate_salary_benchmark(role: str, expected_salary: str, industry: str = "") -> str:
    """生成 HK 薪酬市場參考（DeepSeek）。"""
    ai_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
    industry = industry or "通用"

    prompt = f"""你係 HK 職場薪酬顧問。根據香港 2026 年勞動市場，為以下職位提供薪酬參考：

職位：{role}
行業：{industry}
用戶期望月薪：HK${expected_salary}

請輸出：
1. 該職位 Junior / Mid / Senior 三個等級嘅市場薪酬範圍（HK$/月）
2. 用戶期望係偏低 / 合理 / 偏高（對比 Mid level）
3. 談判建議：開價策略（例：開幾多，最終落幾多）
4. 如果偏低，建議提高到幾多

格式簡潔，用繁體中文，唔超過 150 字。"""

    try:
        resp = ai_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=500,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"⚠️ 生成失敗：{e}"


# ── ATS Match Score ───────────────────────────────────────────────

def calculate_ats_score(jd_text: str, cv_text: str) -> dict:
    """抽取 JD 關鍵詞（DeepSeek），對比 CV 文本（本地比對），計算 ATS Match Score。"""
    ai_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

    prompt = f"""從以下 JD 抽出 15-20 個最重要嘅關鍵詞（skills, tools, qualifications, job-specific terms）。
只輸出詞語清單，逗號分隔，唔加解釋。

JD:
{(jd_text or "")[:2500]}"""

    try:
        resp = ai_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=300,
        )
        raw = resp.choices[0].message.content
    except Exception as e:
        return {"score": 0, "matched": [], "missing": [], "improvement": f"⚠️ 抽取關鍵詞失敗：{e}"}

    keywords = [k.strip() for k in raw.replace("\n", ",").split(",") if k.strip()]
    cv_lower = (cv_text or "").lower()

    matched = [k for k in keywords if k.lower() in cv_lower]
    missing = [k for k in keywords if k not in matched]
    score = round(len(matched) / len(keywords) * 100) if keywords else 0

    if missing:
        improvement = f"建議喺 CV 加入：{', '.join(missing[:3])}"
    else:
        improvement = "CV 已涵蓋大部分關鍵詞！"

    return {"score": score, "matched": matched, "missing": missing, "improvement": improvement}


def format_ats_message(ats: dict, cv_health_score: int = None) -> str:
    """將 calculate_ats_score() 嘅 dict 結果砌成 Telegram 訊息，附帶對比 CV Health baseline 嘅 delta。"""
    score = ats.get("score", 0)
    delta_text = ""
    if cv_health_score is not None:
        delta = score - cv_health_score
        delta_text = f"（比你原本 CV 高 +{delta} 分 ✅）" if delta > 0 else "（需改善 ⚠️）"

    matched = ats.get("matched", [])
    missing = ats.get("missing", [])

    lines = [f"📊 ATS Match Score：{score}/100{delta_text}", ""]
    if matched:
        lines.append(f"✅ 匹配關鍵詞（{len(matched)} 個）：{', '.join(matched)}")
    if len(missing) >= 3:
        lines.append(f"⚠️ 缺漏關鍵詞：{', '.join(missing)}")
    lines.append("")
    if ats.get("improvement"):
        lines.append(f"💡 {ats['improvement']}")
    return "\n".join(lines)


# ── 薪酬談判 Role-play ────────────────────────────────────────────

def generate_negotiate_response(offer_details: str, user_message: str, round_num: int, history: list = None) -> str:
    """薪酬談判 role-play：HR 角色自然回應 + 即時技巧評分。history 係 [{"user":.., "hr":..}, ...]。"""
    ai_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
    history = history or []

    history_text = ""
    if history:
        lines = []
        for h in history:
            if h.get("user"): lines.append(f"候選人：{h['user']}")
            if h.get("hr"):   lines.append(f"HR：{h['hr']}")
        history_text = "【之前對話】\n" + "\n".join(lines) + "\n\n"

    prompt = f"""你係一間香港公司嘅 HR Manager，正在同候選人進行薪酬談判。
Offer details：{offer_details}
你嘅立場：維護公司利益，但可以在合理範圍內加薪最多 10-15%。
第 {round_num} 回合。

{history_text}候選人講：{user_message}

請：
1. 以 HR 身份自然回應（廣東話夾英文）
2. 回應後用分隔線「━━━━━━━━━━」，評分候選人呢句說話：
   - 技巧評分：X/10
   - 優點：[1句]
   - 改善：[1句，如有]"""

    try:
        resp = ai_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.8,
            max_tokens=600,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"⚠️ 回應失敗：{e}"


def extract_negotiate_reply(ai_text: str) -> str:
    """從 generate_negotiate_response() 嘅輸出中抽取淨 HR 回應部分（去除評分），用於下一回合嘅 history context。"""
    return (ai_text or "").split("━━━━━━━━━━")[0].strip()


def generate_negotiate_summary(history: list) -> str:
    """談判結束後嘅總結評分。history 係 [{"user":.., "hr":..}, ...]。"""
    ai_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

    lines = []
    for h in history or []:
        if h.get("user"): lines.append(f"候選人：{h['user']}")
        if h.get("hr"):   lines.append(f"HR：{h['hr']}")
    transcript = "\n".join(lines) or "（無對話記錄）"

    prompt = f"""你係資深薪酬談判教練，根據以下完整談判對話，給出總結評分。

【談判對話記錄】
{transcript}

【輸出格式——廣東話口語】

💼 談判總結
整體表現：[X/10]

✅ 做得好：[2-3點]
⚠️ 下次注意：[1-2點]
💡 最強嘅一句：「[引用候選人說過嘅最佳回應]」
建議最終 counter-offer：HK$[XXK]"""

    try:
        resp = ai_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.6,
            max_tokens=600,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"⚠️ 總結失敗：{e}"


# ── 面試後覆盤分析 ────────────────────────────────────────────────

def generate_debrief(job_info: dict, debrief_text: str) -> str:
    """根據用戶描述嘅真實面試過程，提供專業覆盤分析。job_info 可為 None（未連結特定工）。"""
    ai_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
    job_info = job_info or {}
    job_title = job_info.get("role", "")
    company   = job_info.get("company", "")
    job_line  = f"{job_title}（{company}）" if job_title else "（未指定職位）"

    prompt = f"""你係一位香港職場面試教練。根據以下面試描述，提供專業分析。

職位：{job_line}
面試描述：{debrief_text}

請輸出：
1. 整體表現評級：A / B+ / B / C（附 1 句理由）
2. 強項（✅）：2-3 點
3. 改善點（⚠️）：2-3 點，每點附具體改善方法
4. 最值得注意嘅一個 moment：正面或需改善
5. 下次策略：2 句

用繁體中文，語氣直接務實。"""

    try:
        resp = ai_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.6,
            max_tokens=900,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"⚠️ 分析失敗：{e}"
